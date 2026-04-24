from __future__ import annotations

from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple
from xml.sax.saxutils import escape as _xml_escape

from semi_beam.services.docx_spanish import localize_docx_to_spanish_ar, normalize_spanish_text

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn
    from docx.shared import Mm, Pt, RGBColor
except Exception:  # pragma: no cover
    Document = None
    WD_ORIENT = None
    WD_ALIGN_VERTICAL = None
    WD_TABLE_ALIGNMENT = None
    WD_ALIGN_PARAGRAPH = None
    OxmlElement = None
    parse_xml = None
    nsdecls = None
    qn = None
    Mm = None
    Pt = None
    RGBColor = None


DOCX_A4_WIDTH_MM = 210.0
DOCX_A4_HEIGHT_MM = 297.0
DOCX_MARGIN_MM = 25.4
DOCX_USABLE_WIDTH_MM = DOCX_A4_WIDTH_MM - (2.0 * DOCX_MARGIN_MM)
DOCX_CARD_COLUMNS = 2
DOCX_CARD_CELL_WIDTH_MM = DOCX_USABLE_WIDTH_MM / DOCX_CARD_COLUMNS
DOCX_CARD_IMAGE_WIDTH_MM = 62.0


@dataclass
class MemoriaHeader:
    titulo: str
    cliente_proyecto: str = ""
    autor: str = ""
    fecha: Optional[datetime] = None
    revision: str = ""


@dataclass
class MemoriaCaso:
    unidad: str
    L_carrozable_mm: float
    L_viga_total_mm: float
    descripcion_config: str
    apoyos: List[Tuple[str, str]] = field(default_factory=list)
    cargas: List[Tuple[str, str]] = field(default_factory=list)


@dataclass
class MemoriaResultados:
    q_user_kgmm: float
    x_t_mm: float
    x_d_mm: Optional[float]
    residual_Fy: float
    residual_M0: float
    extremos_V: List[Tuple[str, float, float]] = field(default_factory=list)
    extremos_M: List[Tuple[str, float, float]] = field(default_factory=list)
    vmin_mm: Optional[float] = None
    x_vmin_mm: Optional[float] = None
    utilized_mm: Optional[float] = None
    allowable_mm: Optional[float] = None
    deflection_ok: Optional[bool] = None
    i_source: str = ""


@dataclass
class MemoriaSeccion:
    materiales: List[Tuple[str, str]] = field(default_factory=list)
    fs_min: float = 0.0
    n_vigas: int = 1
    parametros: List[Tuple[str, str]] = field(default_factory=list)
    tabla: List[Sequence[str]] = field(default_factory=list)


def _ensure_docx_lib() -> None:
    if Document is None:
        raise RuntimeError("Falta dependencia 'python-docx'. Instale requirements.txt y reintente.")


def _fmt_num(value: Any, decimals: int = 2) -> str:
    try:
        number = float(value)
    except Exception:
        return normalize_spanish_text(str(value))
    text = f"{number:.{int(decimals)}f}"
    if "." in text:
        text = text.rstrip("0").rstrip(".")
    return text.replace(".", ",")


def _safe_float(value: Any) -> Optional[float]:
    try:
        text = str(value if value is not None else "").strip().replace(",", ".")
        return None if text == "" or text == "-" else float(text)
    except Exception:
        return None


def _clean_text(value: Any) -> str:
    return normalize_spanish_text(str(value if value is not None else ""))


def _append_omml_equation(paragraph, text: str) -> None:
    _ensure_docx_lib()
    expr = _xml_escape(normalize_spanish_text(text or "").strip()) or "0"
    omml = parse_xml(
        f'<m:oMath {nsdecls("m", "w")}>'
        f'<m:r><w:rPr><w:rStyle w:val="Equation"/></w:rPr><m:t>{expr}</m:t></m:r>'
        f"</m:oMath>"
    )
    paragraph._p.append(omml)


def _add_equation_paragraph(doc, text: str):
    p = doc.add_paragraph("")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(3)
    _append_omml_equation(p, text)
    return p


def _add_paragraph(
    doc,
    text: str,
    *,
    bold: bool = False,
    font_size: float = 10.0,
    align=None,
    color: Optional[str] = None,
    spacing_after_pt: float = 6.0,
):
    p = doc.add_paragraph("")
    p.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(spacing_after_pt)
    pf.line_spacing = 1.0
    run = p.add_run(_clean_text(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    font_size: float = 9.0,
    align=None,
    color: Optional[str] = None,
) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    run = p.add_run(_clean_text(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _add_kv_table(doc, rows: Sequence[Tuple[str, str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    widths = [52.0, DOCX_USABLE_WIDTH_MM - 52.0]
    for row_idx, (label, value) in enumerate(rows):
        label_cell = table.rows[row_idx].cells[0]
        value_cell = table.rows[row_idx].cells[1]
        label_cell.width = Mm(widths[0])
        value_cell.width = Mm(widths[1])
        label_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        value_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_shading(label_cell, "E2E8F0")
        _set_cell_shading(value_cell, "F8FAFC")
        _set_cell_text(label_cell, label, bold=True, font_size=8.5, align=WD_ALIGN_PARAGRAPH.LEFT, color="1E293B")
        _set_cell_text(value_cell, value, font_size=8.5, align=WD_ALIGN_PARAGRAPH.LEFT, color="111827")
    doc.add_paragraph("")


def _configure_document(doc) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(DOCX_A4_WIDTH_MM)
    section.page_height = Mm(DOCX_A4_HEIGHT_MM)
    section.left_margin = Mm(DOCX_MARGIN_MM)
    section.right_margin = Mm(DOCX_MARGIN_MM)
    section.top_margin = Mm(DOCX_MARGIN_MM)
    section.bottom_margin = Mm(DOCX_MARGIN_MM)


def _append_image(doc, title: str, path: str, *, width_mm: float = 150.0) -> bool:
    if not path or not Path(path).exists():
        return False
    _add_paragraph(doc, title, bold=True, font_size=9.5, color="334155", spacing_after_pt=3)
    p = doc.add_paragraph("")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Mm(width_mm))
    doc.add_paragraph("")
    return True


def _card_title(card: Dict[str, Any]) -> str:
    if card.get("title"):
        return str(card["title"])
    x_text = card.get("x_mm") or "-"
    return (
        f"Seccion {card.get('sec', '-')} | x = {x_text} mm\n"
        f"h_viga = {_fmt_num(card.get('h_web_mm', 0), 0)} mm | "
        f"tw = {card.get('t_web_in', '-')} in | FS = {card.get('fs_text', '-')}"
    )


def _add_header(doc, header: Optional[MemoriaHeader], caso: Optional[MemoriaCaso]) -> None:
    _add_paragraph(
        doc,
        "Verificacion de viga",
        bold=True,
        font_size=16,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color="1F3A5F",
        spacing_after_pt=3,
    )
    _add_paragraph(
        doc,
        "Memoria de Cálculo y reporte de secciones analizadas",
        font_size=10,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color="4B5563",
        spacing_after_pt=10,
    )
    if header is None and caso is None:
        return
    now = header.fecha if header and header.fecha else datetime.now()
    rows = [
        ("Documento", header.titulo if header else "Memoria de cálculo"),
        ("Cliente / proyecto", header.cliente_proyecto if header else ""),
        ("Autor", header.autor if header else ""),
        ("Revisión", header.revision if header else ""),
        ("Unidad", caso.unidad if caso else ""),
        ("Fecha", now.strftime("%Y-%m-%d %H:%M")),
    ]
    _add_kv_table(doc, [(k, v) for k, v in rows if str(v or "").strip()])


def _add_case_summary(
    doc,
    caso: Optional[MemoriaCaso],
    resultados: Optional[MemoriaResultados],
    seccion: Optional[MemoriaSeccion],
    extras: Dict[str, Any],
) -> None:
    if caso is None:
        return
    _add_paragraph(doc, "0. Resumen del modelo y solicitaciones", bold=True, font_size=12, color="1F3A5F", spacing_after_pt=4)
    _add_kv_table(
        doc,
        [
            ("Configuración", caso.descripcion_config),
            ("Largo carrozable", f"{_fmt_num(caso.L_carrozable_mm, 0)} mm"),
            ("Largo de viga", f"{_fmt_num(caso.L_viga_total_mm, 0)} mm"),
            ("Distancia perno de enganche", f"{_fmt_num(extras.get('dist_perno_mm', 0.0), 0)} mm"),
            ("Peso 1º eje completo", f"{_fmt_num(extras.get('peso_eje1_kg', 0.0), 2)} kg"),
            ("Peso 2º eje completo", f"{_fmt_num(extras.get('peso_eje2_kg', 0.0), 2)} kg"),
            ("Momento flector máximo", f"{_fmt_num(extras.get('mmax_kgcm', 0.0), 2)} kg·cm a {_fmt_num(extras.get('mmax_x_mm', 0.0), 0)} mm"),
        ],
    )
    if caso.apoyos:
        _add_paragraph(doc, "Apoyos y reacciones", bold=True, font_size=9.5, color="334155", spacing_after_pt=3)
        _add_kv_table(doc, [(str(k), str(v)) for k, v in caso.apoyos])
    if caso.cargas:
        _add_paragraph(doc, "Cargas consideradas", bold=True, font_size=9.5, color="334155", spacing_after_pt=3)
        _add_kv_table(doc, [(str(k), str(v)) for k, v in caso.cargas])
    if resultados is not None:
        rows = [
            ("Carga distribuida resuelta q", f"{_fmt_num(resultados.q_user_kgmm, 6)} kg/mm"),
            ("x_t", f"{_fmt_num(resultados.x_t_mm, 0)} mm"),
            ("x_d", "-" if resultados.x_d_mm is None else f"{_fmt_num(resultados.x_d_mm, 0)} mm"),
            ("Residual ΣFy", _fmt_num(resultados.residual_Fy, 6)),
            ("Residual ΣM0", _fmt_num(resultados.residual_M0, 6)),
        ]
        _add_paragraph(doc, "Resultados de equilibrio", bold=True, font_size=9.5, color="334155", spacing_after_pt=3)
        _add_kv_table(doc, rows)
    if seccion is not None:
        rows = list(seccion.materiales)
        rows.extend(seccion.parametros)
        rows.extend(
            [
                ("FS mínimo requerido", _fmt_num(seccion.fs_min, 2)),
                ("Cantidad de vigas", str(seccion.n_vigas)),
            ]
        )
        _add_paragraph(doc, "Datos globales de sección", bold=True, font_size=9.5, color="334155", spacing_after_pt=3)
        _add_kv_table(doc, [(str(k), str(v)) for k, v in rows if str(v or "").strip()])


def _add_global_images(doc, imagenes: Dict[str, str]) -> None:
    ordered = [
        ("fbd", "Diagrama de cuerpo libre"),
        ("v", "Diagrama de corte V(x)"),
        ("m", "Diagrama de momento M(x)"),
        ("deflection", "Deformada"),
        ("stab_long", "Estabilidad longitudinal"),
        ("stab_lat", "Estabilidad lateral"),
        ("secciones", "Tabla de secciones"),
    ]
    inserted = False
    for key, title in ordered:
        inserted = _append_image(doc, title, imagenes.get(key, ""), width_mm=155.0) or inserted
    if inserted:
        doc.add_page_break()


def _add_theory_section(doc) -> None:
    _add_paragraph(doc, "1. Base teórica y metodología aplicada", bold=True, font_size=12, color="1F3A5F", spacing_after_pt=4)
    _add_paragraph(
        doc,
        "La verificación se desarrolla para una sección doble T idealizada como tres rectángulos: "
        "planchuela inferior, alma y planchuela superior. El análisis es elástico lineal y considera "
        "flexión simple en el eje fuerte de la sección.",
        font_size=9,
        color="111827",
        spacing_after_pt=4,
    )
    _add_paragraph(doc, "Metodología aplicada:", bold=True, font_size=9.5, color="334155", spacing_after_pt=2)
    for bullet in [
        "Se determina la geometría total, el eje neutro y el momento de inercia mediante el teorema de ejes paralelos.",
        "Para n vigas iguales en paralelo se adopta I_x,total = n · I_x,1 viga.",
        "La demanda resistente se obtiene a partir del momento flector actuante en cada sección evaluada.",
        "Se verifican flexión, cortante y flecha con el menor factor de seguridad como condición gobernante.",
    ]:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Mm(5)
        p.add_run(_clean_text(f"- {bullet}")).font.size = Pt(8.8)
    _add_paragraph(doc, "Ecuaciones generales de verificación", bold=True, font_size=9.5, color="334155", spacing_after_pt=3)
    for eq in [
        "H = t_inf + h_web + t_sup",
        "y_bar = sum(A_i * y_i) / sum(A_i)",
        "I_x = sum(I_x,i + A_i * d_i^2)",
        "W_crit = I_x / c_max",
        "sigma = M * c / I_x",
        "tau = V * Q / (I_x * t)",
        "W_req = max(M / sigma_adm,sup ; M / sigma_adm,inf)",
        "FS = min(sigma_adm,sup / sigma_sup ; sigma_adm,inf / sigma_inf)",
        "v_utilizada = convexidad - v_min ; v_utilizada <= v_adm",
    ]:
        _add_equation_paragraph(doc, eq)
    doc.add_paragraph("")


def _add_row_memory(doc, card: Dict[str, Any], fs_required: float) -> None:
    sec = str(card.get("sec", "-"))
    _add_paragraph(doc, f"2.{sec} Sección {sec}", bold=True, font_size=10.5, color="334155", spacing_after_pt=4)
    _add_kv_table(
        doc,
        [
            ("Posición x", f"{card.get('x_mm') or '-'} mm"),
            ("Momento actuante M", f"{_fmt_num(card.get('moment_kgcm'), 2)} kg·cm" if card.get("moment_kgcm") is not None else "-"),
            ("Corte actuante V", f"{_fmt_num((card.get('shear') or {}).get('V_kg'), 2)} kg" if (card.get("shear") or {}).get("V_kg") is not None else "-"),
            ("Base ala", f"{_fmt_num(card.get('b_f_mm'), 2)} mm"),
            ("Espesor ala superior", f"{_fmt_num(card.get('t_top_mm'), 2)} mm ({_fmt_num(card.get('t_top_in'), 4)} in)"),
            ("Espesor ala inferior", f"{_fmt_num(card.get('t_bot_mm'), 2)} mm ({_fmt_num(card.get('t_bot_in'), 4)} in)"),
            ("Espesor de alma", f"{_fmt_num(card.get('t_web_mm'), 2)} mm ({card.get('t_web_in', '-')})"),
            ("Altura libre del alma", f"{_fmt_num(card.get('h_web_mm'), 2)} mm"),
            ("σ admisible superior", f"{_fmt_num(card.get('sigma_top_adm_kgcm2'), 2)} kg/cm²" if card.get("sigma_top_adm_kgcm2") is not None else "-"),
            ("σ admisible inferior", f"{_fmt_num(card.get('sigma_bot_adm_kgcm2'), 2)} kg/cm²" if card.get("sigma_bot_adm_kgcm2") is not None else "-"),
        ],
    )
    _add_paragraph(doc, "Desarrollo geométrico y resistente", bold=True, font_size=9.3, color="334155", spacing_after_pt=3)
    for eq in [
        f"H = t_inf + h_web + t_sup = {_fmt_num(card.get('t_bot_mm'), 2)} + {_fmt_num(card.get('h_web_mm'), 2)} + {_fmt_num(card.get('t_top_mm'), 2)} = {_fmt_num(card.get('h_total_mm'), 2)} mm",
        f"y_bar = sum(A_i * y_i) / sum(A_i) = {_fmt_num(card.get('ybar_cm'), 4)} cm",
        f"I_x,1 viga = {_fmt_num(card.get('ix_single_cm4'), 4)} cm^4",
        f"I_x,total = n * I_x,1 viga = {_fmt_num(card.get('ix_total_cm4'), 4)} cm^4",
        f"W_crit = I_x,total / c_max = {_fmt_num(card.get('wcrit_cm3'), 4)} cm^3",
    ]:
        _add_equation_paragraph(doc, eq)
    if card.get("moment_kgcm") is not None:
        _add_paragraph(doc, "Verificación a flexión", bold=True, font_size=9.3, color="334155", spacing_after_pt=3)
        result_fs = _safe_float(getattr(card.get("result"), "FS", None))
        result_wreq = _safe_float(getattr(card.get("result"), "Wreq_cm3", None))
        if result_fs is None:
            fs_values = [v for v in (_safe_float(card.get("fs_top")), _safe_float(card.get("fs_bot"))) if v is not None]
            result_fs = min(fs_values) if fs_values else 0.0
        if result_wreq is None:
            wreq_values = [v for v in (_safe_float(card.get("wreq_top_cm3")), _safe_float(card.get("wreq_bot_cm3"))) if v is not None]
            result_wreq = max(wreq_values) if wreq_values else 0.0
        for eq in [
            f"sigma_sup = M * c_sup / I_x,total = {_fmt_num(card.get('sigma_top_kgcm2'), 4)} kg/cm^2",
            f"sigma_inf = M * c_inf / I_x,total = {_fmt_num(card.get('sigma_bot_kgcm2'), 4)} kg/cm^2",
            f"W_req = {_fmt_num(result_wreq, 4)} cm^3",
            f"FS = {_fmt_num(result_fs, 4)}",
        ]:
            _add_equation_paragraph(doc, eq)
        compliance = float(result_fs or 0.0) >= float(fs_required)
        _add_paragraph(
            doc,
            f"Conclusión: la sección {sec} {'CUMPLE' if compliance else 'NO CUMPLE'} la verificación a flexión respecto del FS mínimo exigido de {_fmt_num(fs_required, 2)}.",
            bold=True,
            font_size=9.2,
            color="0A7F2E" if compliance else "B00020",
            spacing_after_pt=7,
        )
    shear = card.get("shear") or {}
    if shear.get("V_kg") is not None and shear.get("tau_adm_kgcm2") is not None and shear.get("tau_max_kgcm2") is not None:
        _add_paragraph(doc, "Verificación al cortante", bold=True, font_size=9.3, color="334155", spacing_after_pt=3)
        for eq in [
            f"tau = V * Q / (I_x,total * t) = {_fmt_num(shear.get('tau_max_kgcm2'), 4)} kg/cm^2",
            f"tau_adm = sigma_adm,alma / sqrt(3) = {_fmt_num(shear.get('tau_adm_kgcm2'), 4)} kg/cm^2",
            f"FS_cortante = tau_adm / tau = {_fmt_num(shear.get('fs_shear'), 4)}",
        ]:
            _add_equation_paragraph(doc, eq)
        shear_ok = float(shear.get("fs_shear") or 0.0) >= 1.0
        _add_paragraph(
            doc,
            f"Conclusión parcial por cortante: la sección {sec} {'CUMPLE' if shear_ok else 'NO CUMPLE'} en la zona de {shear.get('zone', '-')}.",
            bold=True,
            font_size=8.9,
            color="0A7F2E" if shear_ok else "B00020",
            spacing_after_pt=4,
        )


def _add_deflection_section(doc, resultados: Optional[MemoriaResultados], deflection_context: Optional[Dict[str, Any]]) -> None:
    _add_paragraph(doc, "2.99 Verificación de flecha", bold=True, font_size=10.5, color="334155", spacing_after_pt=4)
    ctx = dict(deflection_context or {})
    if resultados is not None and resultados.vmin_mm is not None:
        ctx.setdefault("vmin_mm", resultados.vmin_mm)
        ctx.setdefault("x_vmin_mm", resultados.x_vmin_mm)
        ctx.setdefault("utilized_mm", resultados.utilized_mm)
        ctx.setdefault("allowable_mm", resultados.allowable_mm)
        ctx.setdefault("ok", resultados.deflection_ok)
        ctx.setdefault("i_source", resultados.i_source)
    if not ctx:
        _add_paragraph(
            doc,
            "No hay resultados de deformada disponibles para documentar la comparación formal de flecha en este reporte.",
            font_size=8.9,
            color="B00020",
            spacing_after_pt=6,
        )
        return
    _add_kv_table(
        doc,
        [
            ("Fuente de inercia", str(ctx.get("i_source", "-") or "-")),
            ("Convexidad adoptada", f"{_fmt_num(ctx.get('camber_mid_mm', 0.0), 2)} mm"),
            ("Flecha mínima obtenida", f"{_fmt_num(ctx.get('vmin_mm', 0.0), 2)} mm"),
            ("Posición de flecha mínima", f"{_fmt_num(ctx.get('x_vmin_mm', 0.0), 0)} mm"),
            ("Límite geométrico inferior", f"{_fmt_num(ctx.get('limit_y_mm', 0.0), 2)} mm"),
            ("Flecha utilizada", f"{_fmt_num(ctx.get('utilized_mm', 0.0), 2)} mm"),
            ("Flecha admisible", f"{_fmt_num(ctx.get('allowable_mm', 0.0), 2)} mm"),
        ],
    )
    _add_equation_paragraph(doc, "v_total(x) = v_precamber(x) + v_cargas(x)")
    _add_equation_paragraph(
        doc,
        f"v_utilizada = convexidad - v_min = {_fmt_num(ctx.get('camber_mid_mm', 0.0), 2)} - ({_fmt_num(ctx.get('vmin_mm', 0.0), 2)}) = {_fmt_num(ctx.get('utilized_mm', 0.0), 2)} mm",
    )
    _add_equation_paragraph(
        doc,
        f"Verificación: {_fmt_num(ctx.get('utilized_mm', 0.0), 2)} {'<=' if ctx.get('ok') else '>'} {_fmt_num(ctx.get('allowable_mm', 0.0), 2)}",
    )
    _add_paragraph(
        doc,
        f"Conclusión parcial por flecha: la viga {'CUMPLE' if ctx.get('ok') else 'NO CUMPLE'} el criterio adoptado de deformación.",
        bold=True,
        font_size=9.0,
        color="0A7F2E" if ctx.get("ok") else "B00020",
        spacing_after_pt=7,
    )


def _add_verification_memory(
    doc,
    cards: Sequence[Dict[str, Any]],
    *,
    fs_required: float,
    resultados: Optional[MemoriaResultados],
    deflection_context: Optional[Dict[str, Any]],
) -> None:
    _add_paragraph(doc, "2. Memoria de cálculo de la verificación de sección", bold=True, font_size=12, color="1F3A5F", spacing_after_pt=4)
    _add_paragraph(
        doc,
        "A continuación se presenta el desarrollo numérico para cada sección analizada con datos de entrada, reemplazo en fórmulas, resultados parciales y conclusión de cumplimiento.",
        font_size=9,
        color="111827",
        spacing_after_pt=6,
    )
    for card in cards:
        _add_row_memory(doc, dict(card), fs_required)
    _add_deflection_section(doc, resultados, deflection_context)


def _add_cards_grid(doc, cards: Sequence[Dict[str, Any]]) -> None:
    cards_with_images = [dict(card) for card in cards if card.get("image_path") and Path(str(card["image_path"])).exists()]
    if not cards_with_images:
        return
    _add_paragraph(doc, "3. Grilla de secciones de viga analizadas", bold=True, font_size=11, color="1F3A5F", spacing_after_pt=6)
    grid = doc.add_table(rows=int((len(cards_with_images) + DOCX_CARD_COLUMNS - 1) / DOCX_CARD_COLUMNS), cols=DOCX_CARD_COLUMNS)
    grid.alignment = WD_TABLE_ALIGNMENT.CENTER
    grid.autofit = False
    grid.style = "Table Grid"
    for idx, card in enumerate(cards_with_images):
        cell = grid.rows[idx // DOCX_CARD_COLUMNS].cells[idx % DOCX_CARD_COLUMNS]
        cell.width = Mm(DOCX_CARD_CELL_WIDTH_MM)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_shading(cell, "F8FAFC")
        _set_cell_text(cell, _card_title(card), bold=True, font_size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER, color="0F172A")
        p = cell.add_paragraph("")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        p.add_run().add_picture(str(card["image_path"]), width=Mm(DOCX_CARD_IMAGE_WIDTH_MM))
    for idx in range(len(cards_with_images), len(grid.rows) * DOCX_CARD_COLUMNS):
        cell = grid.rows[idx // DOCX_CARD_COLUMNS].cells[idx % DOCX_CARD_COLUMNS]
        cell.width = Mm(DOCX_CARD_CELL_WIDTH_MM)
        cell.text = ""
        _set_cell_shading(cell, "FFFFFF")
    doc.add_paragraph("")


def _table_row_groups(headers: Sequence[str]) -> List[Tuple[str, List[int], List[float]]]:
    # Stable column positions from SectionCheckPanel.
    return [
        ("Datos geometricos y solicitacion", [0, 1, 2, 3, 4, 5], [16.0, 20.0, 23.0, 31.0, 14.0, 55.2]),
        ("Propiedades y resultado", [0, 6, 7, 8, 9, 10, 11], [14.0, 24.0, 17.0, 17.0, 28.0, 28.0, 31.2]),
    ]


def _add_verification_tables(
    doc,
    *,
    headers: Sequence[str],
    data: Sequence[Sequence[str]],
    row_ok: Sequence[bool],
    fs_required: float,
    n_beams: int,
) -> None:
    if not headers or not data:
        return
    used_rows = [idx for idx, row in enumerate(data) if any((cell or "").strip() for cell in row)]
    if not used_rows:
        return
    _add_paragraph(
        doc,
        f"4. Tabla de verificacion | FS minimo requerido = {_fmt_num(fs_required, 2)} | Cantidad de vigas = {n_beams}",
        bold=True,
        font_size=11,
        color="1F3A5F",
        spacing_after_pt=6,
    )
    for group_title, col_indices, widths_mm in _table_row_groups(headers):
        available = [(idx, width) for idx, width in zip(col_indices, widths_mm) if idx < len(headers)]
        if not available:
            continue
        _add_paragraph(doc, group_title, bold=True, font_size=9.5, color="334155", spacing_after_pt=4)
        table = doc.add_table(rows=1 + len(used_rows), cols=len(available))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.style = "Table Grid"
        header_row = table.rows[0]
        _set_repeat_header(header_row)
        for col_pos, (src_idx, width_mm) in enumerate(available):
            cell = header_row.cells[col_pos]
            cell.width = Mm(width_mm)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_shading(cell, "1F3A5F")
            _set_cell_text(cell, headers[src_idx], bold=True, font_size=8.0, align=WD_ALIGN_PARAGRAPH.CENTER, color="FFFFFF")
        for row_pos, src_row in enumerate(used_rows, start=1):
            filled = any((cell or "").strip() for cell in data[src_row])
            ok = bool(row_ok[src_row]) if src_row < len(row_ok) else False
            fill = "DCFCE7" if ok else ("FEE2E2" if filled else "F3F4F6")
            for col_pos, (src_idx, width_mm) in enumerate(available):
                cell = table.rows[row_pos].cells[col_pos]
                cell.width = Mm(width_mm)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                _set_cell_shading(cell, fill)
                value = data[src_row][src_idx] if src_idx < len(data[src_row]) else ""
                _set_cell_text(cell, value, font_size=8.0, align=WD_ALIGN_PARAGRAPH.CENTER, color="111827")
        doc.add_paragraph("")


def _verification_from_legacy_inputs(
    seccion: Optional[MemoriaSeccion],
    resultados: Optional[MemoriaResultados],
    extras: Dict[str, Any],
) -> Dict[str, Any]:
    headers = [
        "Sección",
        "x [mm]",
        "h_viga [mm]",
        "Espesor",
        "FS",
        "M [kg·cm]",
        "Jx [cm^4]",
        "ȳ [cm]",
        "c_max [cm]",
        "Wcrit [cm^3]",
        "Wreq [cm^3]",
        "σmax [kg/cm²]",
    ]
    rows: List[List[str]] = []
    row_ok: List[bool] = []
    cards: List[Dict[str, Any]] = []
    fs_required = float(seccion.fs_min) if seccion is not None else 0.0
    n_beams = int(seccion.n_vigas) if seccion is not None else 1
    for idx, row in enumerate(list(extras.get("flex_rows", []))):
        fs = _safe_float(row.get("FS")) or 0.0
        sec = str(row.get("sec") or idx + 1)
        rows.append(
            [
                sec,
                str(row.get("x_mm", "")),
                str(row.get("h_web_mm", "")),
                str(row.get("t_web_in", "")),
                _fmt_num(fs, 2),
                _fmt_num(row.get("M_kgcm", 0.0), 2),
                str(row.get("Jx_cm4", "")),
                str(row.get("ybar_cm", "")),
                str(row.get("cmax_cm", "")),
                _fmt_num(row.get("Wcrit_cm3", 0.0), 2),
                _fmt_num(row.get("Wreq_cm3", 0.0), 2),
                _fmt_num(row.get("sigma_max", 0.0), 2),
            ]
        )
        row_ok.append(bool(fs_required <= 0.0 or fs >= fs_required))
        cards.append(
            {
                "sec": sec,
                "x_mm": row.get("x_mm", ""),
                "h_web_mm": row.get("h_web_mm", 0.0),
                "t_web_in": row.get("t_web_in", ""),
                "fs_text": _fmt_num(fs, 2),
                "moment_kgcm": row.get("M_kgcm", 0.0),
                "wcrit_cm3": row.get("Wcrit_cm3", 0.0),
                "wreq_top_cm3": row.get("Wreq_cm3", 0.0),
                "wreq_bot_cm3": row.get("Wreq_cm3", 0.0),
                "sigma_top_kgcm2": row.get("sigma_max", 0.0),
                "sigma_bot_kgcm2": row.get("sigma_max", 0.0),
                "fs_top": fs,
                "fs_bot": fs,
            }
        )
    deflection_context = None
    if resultados is not None and resultados.vmin_mm is not None:
        deflection_context = {
            "vmin_mm": resultados.vmin_mm,
            "x_vmin_mm": resultados.x_vmin_mm,
            "utilized_mm": resultados.utilized_mm,
            "allowable_mm": resultados.allowable_mm,
            "ok": resultados.deflection_ok,
            "i_source": resultados.i_source,
        }
    return {
        "headers": headers,
        "data": rows,
        "row_ok": row_ok,
        "cards": cards,
        "fs_required": fs_required,
        "n_beams": n_beams,
        "deflection_context": deflection_context,
    }


def ensure_memoria_template(path: str) -> str:
    """Backward-compatible no-op for callers that still pass a template path."""
    p = Path(path)
    p.parent.mkdir(parents=True, exist_ok=True)
    return str(p)


def default_template_path() -> str:
    return str(Path.cwd() / "assets" / "templates" / "memoria_base.docx")


def export_memoria_docx(
    out_docx_path: str,
    *,
    template_path: str = "",
    header: Optional[MemoriaHeader] = None,
    caso: Optional[MemoriaCaso] = None,
    resultados: Optional[MemoriaResultados] = None,
    seccion: Optional[MemoriaSeccion] = None,
    imagenes: Optional[Dict[str, str]] = None,
    extras: Optional[Dict[str, Any]] = None,
    verification: Optional[Dict[str, Any]] = None,
) -> None:
    """Exporta la memoria completa con el formato canónico del verificador de viga.

    `template_path` se conserva solo por compatibilidad: el pipeline actual no usa
    plantillas externas para evitar un segundo renderizador DOCX.
    """
    _ensure_docx_lib()
    extras = dict(extras or {})
    imagenes = dict(imagenes or {})
    verification = dict(verification or _verification_from_legacy_inputs(seccion, resultados, extras))
    cards = list(verification.get("cards", []))
    headers = list(verification.get("headers", []))
    data = list(verification.get("data", []))
    row_ok = list(verification.get("row_ok", []))
    fs_required = float(verification.get("fs_required", seccion.fs_min if seccion else 0.0) or 0.0)
    n_beams = int(verification.get("n_beams", seccion.n_vigas if seccion else 1) or 1)
    deflection_context = verification.get("deflection_context")

    if not cards and not data and caso is None:
        raise RuntimeError("No hay datos de memoria de cálculo para exportar.")

    doc = Document()
    _configure_document(doc)
    _add_header(doc, header, caso)
    _add_case_summary(doc, caso, resultados, seccion, extras)
    _add_global_images(doc, imagenes)
    _add_theory_section(doc)
    _add_verification_memory(
        doc,
        cards,
        fs_required=fs_required,
        resultados=resultados,
        deflection_context=deflection_context,
    )
    _add_cards_grid(doc, cards)
    _add_verification_tables(
        doc,
        headers=headers,
        data=data,
        row_ok=row_ok,
        fs_required=fs_required,
        n_beams=n_beams,
    )

    out = Path(out_docx_path if str(out_docx_path).lower().endswith(".docx") else f"{out_docx_path}.docx")
    out.parent.mkdir(parents=True, exist_ok=True)
    localize_docx_to_spanish_ar(doc)
    doc.save(str(out))
