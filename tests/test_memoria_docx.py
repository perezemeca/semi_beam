import base64
import math
import os
import tempfile
import zipfile
from datetime import datetime

from docx import Document

from semi_beam.services.memoria_calculo_docx import (
    export_memoria_docx,
    MemoriaCaso,
    MemoriaHeader,
    MemoriaResultados,
)


_PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO2p6Z8AAAAASUVORK5CYII="
)


def _write_png(path: str) -> None:
    with open(path, "wb") as fh:
        fh.write(_PNG_1X1)


def _document_xml(path) -> str:
    with zipfile.ZipFile(path) as zf:
        return zf.read("word/document.xml").decode("utf-8")


def _document_text(path) -> str:
    doc = Document(path)
    texts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(p.text for p in cell.paragraphs if p.text)
    return "\n".join(texts)


def _export_with_verification(tmp_path, verification: dict) -> str:
    out = tmp_path / "memoria.docx"
    export_memoria_docx(out, verification=verification)
    return _document_xml(out)


def test_export_memoria_docx_includes_deflection_image():
    with tempfile.TemporaryDirectory() as td:
        out = os.path.join(td, "memoria.docx")

        img_fbd = os.path.join(td, "fbd.png")
        img_v = os.path.join(td, "v.png")
        img_m = os.path.join(td, "m.png")
        img_defl = os.path.join(td, "deflection.png")
        for path in (img_fbd, img_v, img_m, img_defl):
            _write_png(path)

        header = MemoriaHeader(titulo="Memoria Test", fecha=datetime.now())
        caso = MemoriaCaso(
            unidad="Caso 1",
            L_carrozable_mm=1000.0,
            L_viga_total_mm=1200.0,
            descripcion_config="Dummy",
            apoyos=[("Rp1", "x=0 mm; R=0 kg")],
            cargas=[("P1", "x=200 mm; P=100 kg")],
        )
        resultados = MemoriaResultados(
            q_user_kgmm=0.1,
            x_t_mm=500.0,
            x_d_mm=None,
            residual_Fy=0.0,
            residual_M0=0.0,
            extremos_V=[],
            extremos_M=[],
            vmin_mm=-12.0,
            x_vmin_mm=450.0,
            utilized_mm=42.0,
            allowable_mm=60.0,
            deflection_ok=True,
            i_source="tabla de secciones",
        )

        export_memoria_docx(
            out,
            header=header,
            caso=caso,
            resultados=resultados,
            seccion=None,
            imagenes={
                "fbd": img_fbd,
                "v": img_v,
                "m": img_m,
                "deflection": img_defl,
            },
            extras={},
        )

        doc = Document(out)
        assert len(doc.inline_shapes) == 4
        text = "\n".join(p.text for p in doc.paragraphs if p.text)
        assert "Memoria de Cálculo" in text
        with zipfile.ZipFile(out) as zf:
            document_xml = zf.read("word/document.xml").decode("utf-8")
            styles_xml = zf.read("word/styles.xml").decode("utf-8")
            settings_xml = zf.read("word/settings.xml").decode("utf-8")
            core_xml = zf.read("docProps/core.xml").decode("utf-8")
        assert "es-AR" in document_xml
        assert "es-AR" in styles_xml
        assert "es-AR" in settings_xml
        assert "es-AR" in core_xml
        assert "en-US" not in document_xml
        assert "en-US" not in styles_xml


def test_docx_renders_non_verifiable_material_error(tmp_path):
    xml = _export_with_verification(
        tmp_path,
        {
            "fs_required": 1.5,
            "n_beams": 1,
            "cards": [
                {
                    "sec": "1",
                    "fs_text": "ERR MAT",
                    "ok": False,
                    "material_error": True,
                    "missing_material_components": [
                        {"component": "Bastidor lateral", "material": "SAE1010"},
                    ],
                    "component_checks": [
                        {
                            "component": "Bastidor lateral",
                            "material": "SAE1010",
                            "y_inf_cm": 0.0,
                            "y_sup_cm": 17.0,
                            "cmax_cm": 8.5,
                            "sigma_calc_kgcm2": 1200.0,
                            "sigma_adm_kgcm2": None,
                            "fs": None,
                        },
                    ],
                }
            ],
        },
    )

    assert "ERR MAT" in xml
    assert "Sección no verificable" in xml
    assert "MATERIAL FALTANTE" in xml
    assert "FS = 0.0000" not in xml
    assert "FS = 0,0000" not in xml


def test_docx_renders_non_verifiable_chapon_context_error(tmp_path):
    xml = _export_with_verification(
        tmp_path,
        {
            "fs_required": 1.5,
            "n_beams": 1,
            "cards": [
                {
                    "sec": "1",
                    "fs_text": "ERR CHAPÓN",
                    "ok": False,
                    "chapon_context_error": True,
                    "chapon_context_missing_fields": ["largo_viga_mm", "posicion_perno_mm"],
                }
            ],
        },
    )

    assert "ERR CHAPÓN" in xml
    assert "contexto longitudinal" in xml
    assert "largo_viga_mm" in xml
    assert "posicion_perno_mm" in xml
    assert "FS = 0.0000" not in xml
    assert "FS = 0,0000" not in xml


def test_docx_renders_table_values(tmp_path):
    xml = _export_with_verification(
        tmp_path,
        {
            "fs_required": 1.5,
            "n_beams": 1,
            "cards": [
                {
                    "sec": "1",
                    "fs_text": "2.31",
                    "ok": True,
                    "table_values": {
                        "FS": "2.31",
                        "Jx_cm4": "12345.67",
                        "ybar_cm": "12.34",
                        "cmax_cm": "56.78",
                        "Wcrit_cm3": "987.65",
                        "Wreq_cm3": "432.10",
                        "sigma_max_kgcm2": "1500",
                    },
                }
            ],
        },
    )

    assert "Resultado mostrado en tabla" in xml
    assert "2.31" in xml
    assert "12345.67" in xml
    assert "987.65" in xml
    assert "432.10" in xml
    assert "1500" in xml


def test_docx_renders_simple_web_configuration(tmp_path):
    out = tmp_path / "memoria.docx"
    export_memoria_docx(
        out,
        verification={
            "fs_required": 1.5,
            "n_beams": 1,
            "cards": [
                {
                    "sec": "1",
                    "fs_text": "2.31",
                    "ok": True,
                    "web_configuration_label": "Simple",
                    "double_web_enabled": False,
                    "double_web_inner_face_offset_mm": None,
                    "double_web_clear_gap_mm": None,
                    "t_web_mm": 6.35,
                    "t_web_in": "1/4",
                }
            ],
        },
    )

    text = _document_text(out)

    assert "Configuración de alma" in text
    assert "Simple" in text


def test_docx_renders_double_web_configuration(tmp_path):
    out = tmp_path / "memoria.docx"
    export_memoria_docx(
        out,
        verification={
            "fs_required": 1.5,
            "n_beams": 1,
            "cards": [
                {
                    "sec": "1",
                    "fs_text": "2.31",
                    "ok": True,
                    "web_configuration_label": "Doble",
                    "double_web_enabled": True,
                    "double_web_inner_face_offset_mm": 20.0,
                    "double_web_clear_gap_mm": 40.0,
                    "t_web_mm": 6.35,
                    "t_web_in": "1/4",
                }
            ],
        },
    )

    text = _document_text(out)

    assert "Configuración de alma" in text
    assert "Doble" in text
    assert "Distancia centro" in text
    assert "20" in text
    assert "Luz libre" in text
    assert "40" in text
    assert "Espesor de cada alma" in text


def test_docx_renders_double_web_error_without_zero_fs(tmp_path):
    out = tmp_path / "memoria.docx"
    export_memoria_docx(
        out,
        verification={
            "fs_required": 1.5,
            "n_beams": 1,
            "cards": [
                {
                    "sec": "1",
                    "fs_text": "ERR DOBLE ALMA",
                    "ok": False,
                    "web_configuration_label": "Doble",
                    "double_web_enabled": True,
                    "double_web_inner_face_offset_mm": 60.0,
                    "double_web_clear_gap_mm": 120.0,
                    "double_web_error": "La doble alma excede el ancho de la planchuela.",
                    "t_web_mm": 6.35,
                    "t_web_in": "1/4",
                    "table_values": {"FS": "ERR DOBLE ALMA"},
                }
            ],
        },
    )

    xml = _document_xml(out)
    text = _document_text(out)

    assert "ERR DOBLE ALMA" in text
    assert "Sección no verificable" in text
    assert "FS = 0.0000" not in xml
    assert "FS = 0,0000" not in xml


def test_docx_renders_no_flex_demand_without_inf_or_chapon_error(tmp_path):
    no_demand_text = "Sin demanda a flexi\u00f3n"
    out = tmp_path / "memoria.docx"
    export_memoria_docx(
        out,
        verification={
            "fs_required": 1.5,
            "n_beams": 1,
            "cards": [
                {
                    "sec": "1",
                    "x_mm": "1000",
                    "fs_text": no_demand_text,
                    "ok": True,
                    "include_chapon": True,
                    "chapon_context_error": True,
                    "chapon_context_missing_fields": ["largo_viga_mm", "posicion_perno_mm"],
                    "moment_kgcm": 0.0,
                    "table_values": {
                        "FS": no_demand_text,
                        "Jx_cm4": "12345.67",
                        "ybar_cm": "12.34",
                        "cmax_cm": "56.78",
                        "Wcrit_cm3": "987.65",
                        "Wreq_cm3": "0",
                        "sigma_max_kgcm2": "0",
                    },
                    "wreq_top_cm3": 0.0,
                    "wreq_bot_cm3": 0.0,
                    "sigma_top_kgcm2": 0.0,
                    "sigma_bot_kgcm2": 0.0,
                    "fs_top": math.inf,
                    "fs_bot": math.inf,
                    "component_checks": [
                        {
                            "component": "Viga principal - ala superior",
                            "material": "SAE1010",
                            "y_inf_cm": 0.0,
                            "y_sup_cm": 1.0,
                            "cmax_cm": 1.0,
                            "sigma_calc_kgcm2": 0.0,
                            "sigma_adm_kgcm2": 1400.0,
                            "fs": math.inf,
                            "wreq_cm3": 0.0,
                            "no_flex_demand": True,
                        },
                    ],
                }
            ],
        },
    )

    xml = _document_xml(out)
    text = _document_text(out)

    assert no_demand_text in text
    assert "ERR CHAP" not in text
    assert "FS = inf" not in xml
    assert "&gt;inf&lt;" not in xml
    assert "FS = nan" not in xml
    assert "&gt;nan&lt;" not in xml.lower()
    assert "100000000" not in xml


def test_docx_renders_optional_component_status(tmp_path):
    xml = _export_with_verification(
        tmp_path,
        {
            "fs_required": 1.5,
            "n_beams": 1,
            "cards": [
                {
                    "sec": "1",
                    "x_mm": "2500",
                    "fs_text": "2.31",
                    "ok": True,
                    "include_bastidor_lateral": True,
                    "bastidor_lateral_structural": False,
                    "bastidor_lateral_included": False,
                    "include_piso": True,
                    "piso_structural": False,
                    "piso_included": False,
                    "include_chapon": True,
                    "chapon_included": False,
                    "chapon_context_error": False,
                }
            ],
        },
    )

    assert "Estado de componentes opcionales" in xml
    assert "Bastidor lateral" in xml
    assert "no considerado estructural" in xml
    assert "Piso" in xml
    assert "Chapón" in xml
    assert "fuera del tramo" in xml
