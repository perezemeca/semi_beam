import os
import zipfile
from types import SimpleNamespace

os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")

from docx import Document
from PySide6.QtWidgets import QApplication, QLabel, QMessageBox, QPushButton, QScrollArea, QTableWidget, QToolButton

from semi_beam.domain.loads import DistUniform, PointForce, PointMoment
from semi_beam.domain.beam import Beam
from semi_beam.domain.cases import BeamCase
from semi_beam.domain.supports import FixedSupport, TandemSupport
from semi_beam.domain.unknowns import UnknownUniformLoad
from semi_beam.sections.i_section import ISection
from semi_beam.services.study_storage import load_study_file, save_study_file
from semi_beam.ui.main_window import (
    DIRECTIONAL_OFFSET_MAX_MM,
    DIRECTIONAL_OFFSET_MIN_MM,
    FBDApp,
)
from semi_beam.ui.section_check_panel import SectionCheckPanel


def _app():
    app = QApplication.instance()
    if app is None:
        app = QApplication([])
    return app


def test_section_panel_exposes_new_thickness_options():
    _app()
    panel = SectionCheckPanel()

    top_options = {panel.cmb_t_top.itemData(i) for i in range(panel.cmb_t_top.count())}
    web_options = {panel._tweb_widgets[0].itemData(i) for i in range(panel._tweb_widgets[0].count())}

    assert {"5/16", "3/8", "7/16"}.issubset(top_options)
    assert "7/16" in web_options
    assert panel.cmb_t_top.itemText(panel.cmb_t_top.findData("7/16")) == "7/16 - 11.11 mm"
    assert not hasattr(panel, "btn_export_report")
    assert panel.chk_bastidor_lateral.text() == "Agregar bastidor lateral"
    assert panel.chk_bastidor_lateral.isChecked() is False
    assert panel.chk_bastidor_lateral_structural.isChecked() is True
    assert panel.chk_bastidor_lateral_structural.isVisible() is False
    assert panel.n_bastidor_lateral_altura.minimum() == 130.0
    assert panel.n_bastidor_lateral_altura.maximum() == 170.0
    assert panel.n_bastidor_lateral_altura.value() == 170.0
    piso_options = {panel.cmb_espesor_piso.itemData(i) for i in range(panel.cmb_espesor_piso.count())}
    assert panel.chk_piso.text() == "Agregar piso"
    assert panel.chk_piso.isChecked() is False
    assert not hasattr(panel, "chk_piso_structural")
    assert panel.n_piso_width.value() == 2430.0
    assert panel.n_piso_width.isEnabled() is False
    assert {2.0, 3.0, 4.0, 3.175, 4.7625}.issubset(piso_options)
    assert panel.cmb_mat_piso.count() == panel.cmb_mat_top.count()
    chapon_options = {panel.cmb_espesor_chapon.itemData(i) for i in range(panel.cmb_espesor_chapon.count())}
    assert panel.chk_chapon.text() == "Agregar chapón"
    assert panel.chk_chapon.isChecked() is False
    assert {6.35, 7.9375, 9.525}.issubset(chapon_options)


def test_section_panel_lateral_frame_is_always_structural():
    _app()
    panel = SectionCheckPanel()
    panel.set_moment_provider(lambda _x_mm: 125000.0)
    panel.tbl.item(0, panel.COL_X).setText("1000")
    panel.tbl.item(0, panel.COL_HWEB).setText("450")
    panel._recompute_all()
    base_jx = float(panel.tbl.item(0, panel.COL_JX).text())

    panel.chk_bastidor_lateral.setChecked(True)
    panel.chk_bastidor_lateral_structural.setChecked(False)
    panel._recompute_all()
    legacy_false_jx = float(panel.tbl.item(0, panel.COL_JX).text())

    assert panel.chk_bastidor_lateral_structural.isChecked() is True
    assert legacy_false_jx > base_jx


def test_section_panel_floor_included_always_contributes_and_width_changes_properties():
    _app()
    panel = SectionCheckPanel()
    panel.set_moment_provider(lambda _x_mm: 125000.0)
    panel.tbl.item(0, panel.COL_X).setText("1000")
    panel.tbl.item(0, panel.COL_HWEB).setText("450")
    panel._recompute_all()
    base_jx = float(panel.tbl.item(0, panel.COL_JX).text())

    panel.chk_piso.setChecked(True)
    panel.n_piso_width.setValue(1000.0)
    panel._recompute_all()
    narrow_sec = panel._make_section(450.0, 0.25)
    narrow_area = sum(rect.area_mm2 for rect in narrow_sec.rects)
    narrow_jx = float(panel.tbl.item(0, panel.COL_JX).text())

    panel.n_piso_width.setValue(2400.0)
    panel.cmb_espesor_piso.setCurrentIndex(panel.cmb_espesor_piso.findData(4.7625))
    panel._recompute_all()
    wide_sec = panel._make_section(450.0, 0.25)
    wide_area = sum(rect.area_mm2 for rect in wide_sec.rects)
    wide_jx = float(panel.tbl.item(0, panel.COL_JX).text())
    piso_rects = [r for r in wide_sec.rects if r.label == "piso"]

    assert narrow_jx > base_jx
    assert wide_area > narrow_area
    assert wide_jx != narrow_jx
    assert wide_sec.includes_piso is True
    assert len(piso_rects) == 1
    assert piso_rects[0].x0_mm == -1200.0
    assert piso_rects[0].b_mm == 2400.0
    assert piso_rects[0].y0_mm == wide_sec.base_section.H_mm
    assert piso_rects[0].h_mm == 4.7625


def test_section_panel_floor_not_included_does_not_contribute():
    _app()
    panel = SectionCheckPanel()
    panel.set_moment_provider(lambda _x_mm: 125000.0)
    panel.tbl.item(0, panel.COL_X).setText("1000")
    panel.tbl.item(0, panel.COL_HWEB).setText("450")
    panel.n_piso_width.setValue(2400.0)

    sec = panel._make_section(450.0, 0.25)
    panel._recompute_all()

    assert isinstance(sec, ISection)
    assert float(panel.tbl.item(0, panel.COL_JX).text()) > 0.0
    assert all(getattr(rect, "label", "") != "piso" for rect in getattr(sec, "rects", ()))


def test_section_panel_adds_chapon_only_inside_longitudinal_range():
    _app()
    panel = SectionCheckPanel()
    panel.set_moment_provider(lambda _x_mm: 125000.0)
    panel.tbl.item(0, panel.COL_X).setText("1500")
    panel.tbl.item(0, panel.COL_HWEB).setText("450")
    panel.tbl.item(1, panel.COL_X).setText("2500")
    panel.tbl.item(1, panel.COL_HWEB).setText("450")
    panel._recompute_all()
    base_inside_jx = float(panel.tbl.item(0, panel.COL_JX).text())
    base_outside_jx = float(panel.tbl.item(1, panel.COL_JX).text())

    panel.chk_chapon.setChecked(True)
    panel.n_chapon_length.setValue(2000.0)
    panel.cmb_espesor_chapon.setCurrentIndex(panel.cmb_espesor_chapon.findData(9.525))
    panel._recompute_all()
    inside_jx = float(panel.tbl.item(0, panel.COL_JX).text())
    outside_jx = float(panel.tbl.item(1, panel.COL_JX).text())
    inside_sec = panel._make_section(450.0, 0.25, station_mm=1500.0)
    outside_sec = panel._make_section(450.0, 0.25, station_mm=2500.0)
    chapon_rects = [r for r in inside_sec.rects if r.label == "chapon"]

    assert inside_jx > base_inside_jx
    assert outside_jx == base_outside_jx
    assert inside_sec.includes_chapon is True
    assert isinstance(outside_sec, ISection)
    assert len(chapon_rects) == 1
    assert chapon_rects[0].x0_mm == -525.0
    assert chapon_rects[0].b_mm == 1050.0
    assert chapon_rects[0].y0_mm == -9.525
    assert chapon_rects[0].h_mm == 9.525

    panel.n_chapon_length.setValue(3000.0)
    panel._recompute_all()
    extended_sec = panel._make_section(450.0, 0.25, station_mm=2500.0)
    assert extended_sec.includes_chapon is True
    assert extended_sec.chapon_end_mm == 3000.0


def test_section_panel_governs_by_component_when_all_reinforcements_are_active():
    _app()
    panel = SectionCheckPanel()
    panel.set_moment_provider(lambda _x_mm: 250000.0)
    for combo, mat_id in (
        (panel.cmb_mat_top, "STR900MC"),
        (panel.cmb_mat_bot, "STR900MC"),
        (panel.cmb_mat_web, "STR900MC"),
        (panel.cmb_mat_piso, "STR900MC"),
    ):
        idx = combo.findData(mat_id)
        if idx >= 0:
            combo.setCurrentIndex(idx)
    panel.chk_bastidor_lateral.setChecked(True)
    panel.chk_piso.setChecked(True)
    panel.chk_chapon.setChecked(True)
    panel.tbl.item(0, panel.COL_X).setText("1500")
    panel.tbl.item(0, panel.COL_HWEB).setText("450")

    panel._recompute_all()
    card = panel._collect_section_export_cards()[0]
    checks = card["component_checks"]
    governing = card["governing_component"]
    fs_cell = float(panel.tbl.item(0, panel.COL_FS).text())

    components = {(row["component"], row["material"]) for row in checks}
    expected_min_fs = min(float(row["fs"]) for row in checks if row["fs"] is not None)

    assert ("Bastidor lateral", "SAE1010") in components
    assert ("Chapón", "SAE1010") in components
    assert ("Piso", "STR900MC") in components
    assert governing is not None
    assert abs(float(governing["fs"]) - expected_min_fs) < 1e-9
    assert abs(fs_cell - expected_min_fs) < 0.02


def test_section_panel_exports_verification_report(tmp_path):
    _app()
    panel = SectionCheckPanel()
    panel.set_moment_provider(lambda _x_mm: 125000.0)
    panel.set_shear_provider(lambda _x_mm: 4200.0)
    panel.set_deflection_context(
        SimpleNamespace(
            vmin_mm=-12.0,
            x_vmin_mm=850.0,
            utilized_mm=42.0,
            allowable_mm=60.0,
            limit_y_mm=-30.0,
            camber_mid_mm=30.0,
            ok=True,
        ),
        i_source="tabla de secciones",
    )
    panel.tbl.item(0, panel.COL_X).setText("1000")
    panel.tbl.item(0, panel.COL_HWEB).setText("450")
    panel._recompute_all()

    out = tmp_path / "memoria_calculo.docx"
    panel.export_verification_report(str(out))

    assert out.exists()
    assert out.stat().st_size > 0
    doc = Document(out)
    assert len(doc.inline_shapes) == 1
    assert len(doc.tables) >= 3
    assert doc.paragraphs[0].text == "Verificación de viga"
    all_text = "\n".join(p.text for p in doc.paragraphs if p.text)
    assert "Base teórica y metodología aplicada" in all_text
    assert "Memoria de cálculo de la verificación de sección" in all_text
    assert "Verificación al cortante" in all_text
    assert "Verificación de flecha" in all_text
    assert round(doc.sections[0].page_width.mm) == 210
    assert round(doc.sections[0].page_height.mm) == 297
    assert round(doc.sections[0].left_margin.mm, 1) == 25.4
    with zipfile.ZipFile(out) as zf:
        document_xml = zf.read("word/document.xml").decode("utf-8")
        styles_xml = zf.read("word/styles.xml").decode("utf-8")
        settings_xml = zf.read("word/settings.xml").decode("utf-8")
        core_xml = zf.read("docProps/core.xml").decode("utf-8")
    assert "es-AR" in document_xml
    assert "es-AR" in styles_xml
    assert "es-AR" in settings_xml
    assert "es-AR" in core_xml
    assert "en-US" not in document_xml
    assert "<m:oMath" in document_xml


def test_study_file_roundtrip_restores_ui_state(tmp_path):
    _app()
    window = FBDApp()
    window.tabs.setCurrentWidget(window.tab_semi)

    semi = window.tab_semi
    semi.cmb_semi_tipo.setCurrentIndex(1)
    semi.cmb_config.setCurrentIndex(2)
    semi.Lc.setValue(13600.0)
    semi.x_front_or_kp.setValue(1250.0)
    semi.R_front_or_kp.setValue(9400.0)
    semi.Rt.setValue(22200.0)
    semi.chk_show_deflection.setChecked(False)
    semi._add_load_row(load_type="Puntual")
    semi.tbl.item(0, semi.COL_MAG).setText("5000")
    semi.tbl.item(0, semi.COL_POS).setText("2500")
    semi.section_panel.cmb_t_top.setCurrentIndex(semi.section_panel.cmb_t_top.findData("7/16"))
    semi.section_panel.chk_bastidor_lateral.setChecked(True)
    semi.section_panel.n_bastidor_lateral_altura.setValue(150.0)
    semi.section_panel.chk_piso.setChecked(True)
    semi.section_panel.n_piso_width.setValue(1800.0)
    semi.section_panel.cmb_espesor_piso.setCurrentIndex(semi.section_panel.cmb_espesor_piso.findData(3.175))
    semi.section_panel.cmb_mat_piso.setCurrentIndex(semi.section_panel.cmb_mat_piso.findData("F24"))
    semi.section_panel.chk_chapon.setChecked(True)
    semi.section_panel.n_chapon_length.setValue(2750.0)
    semi.section_panel.cmb_espesor_chapon.setCurrentIndex(semi.section_panel.cmb_espesor_chapon.findData(7.9375))
    semi.section_panel._double_web_widgets[0].setChecked(True)
    semi.section_panel.tbl.item(0, semi.section_panel.COL_DOUBLE_WEB_OFFSET).setText("20")
    semi.section_panel.tbl.item(0, semi.section_panel.COL_X).setText("1000")
    semi.section_panel.tbl.item(0, semi.section_panel.COL_HWEB).setText("450")

    reactions = window.tab_reactions
    reactions.mode.setCurrentIndex(1)
    reactions.offset.setValue(3700.0)

    payload = window._export_study_state()
    study_path = tmp_path / "demo.sbeam"
    save_study_file(study_path, payload)

    restored = FBDApp()
    restored._apply_study_state(load_study_file(study_path))

    restored_semi = restored.tab_semi
    restored_reactions = restored.tab_reactions

    assert restored.tabs.currentIndex() == 1
    assert restored_semi.cmb_semi_tipo.currentIndex() == 1
    assert restored_semi.cmb_config.currentText() == semi.cmb_config.currentText()
    assert restored_semi.Lc.value() == 13600.0
    assert restored_semi.x_front_or_kp.value() == 1250.0
    assert restored_semi.chk_show_deflection.isChecked() is True
    assert restored_semi.deflection_enabled() is True
    assert restored_semi._type_combo(0).currentText() == "Puntual"
    assert restored_semi.tbl.item(0, restored_semi.COL_POS).text() == "2500"
    assert restored_semi.tbl.item(0, restored_semi.COL_MAG).text() == "5000"
    assert restored_semi.section_panel.cmb_t_top.currentData() == "7/16"
    assert restored_semi.section_panel.chk_bastidor_lateral.isChecked() is True
    assert restored_semi.section_panel.chk_bastidor_lateral_structural.isChecked() is True
    assert restored_semi.section_panel.n_bastidor_lateral_altura.value() == 150.0
    assert restored_semi.section_panel.chk_piso.isChecked() is True
    assert restored_semi.section_panel.n_piso_width.value() == 1800.0
    assert restored_semi.section_panel.cmb_espesor_piso.currentData() == 3.175
    assert restored_semi.section_panel.cmb_mat_piso.currentData() == "F24"
    assert restored_semi.section_panel.chk_chapon.isChecked() is True
    assert restored_semi.section_panel.n_chapon_length.value() == 2750.0
    assert restored_semi.section_panel.cmb_espesor_chapon.currentData() == 7.9375
    assert restored_semi.section_panel._row_double_web_enabled(0) is True
    assert restored_semi.section_panel._row_double_web_offset_mm(0) == 20.0
    assert restored_semi.section_panel.tbl.item(0, restored_semi.section_panel.COL_X).text() == "1000"
    assert restored_semi.section_panel.tbl.item(0, restored_semi.section_panel.COL_HWEB).text() == "450"
    assert restored_reactions.mode.currentIndex() == 1
    assert restored_reactions.offset.value() == 3700.0


def test_axis_lift_state_roundtrip_and_old_default(tmp_path):
    _app()
    window = FBDApp()
    semi = window.tab_semi
    semi.chk_axis_lift.setChecked(True)

    payload = window._export_study_state()
    assert payload["tabs"]["semirremolque"]["axis_lift_enabled"] is True

    study_path = tmp_path / "axis_lift.sbeam"
    save_study_file(study_path, payload)

    restored = FBDApp()
    restored._apply_study_state(load_study_file(study_path))
    assert restored.tab_semi.chk_axis_lift.isChecked() is True

    old_payload = window._export_study_state()
    del old_payload["tabs"]["semirremolque"]["axis_lift_enabled"]
    restored_old = FBDApp()
    restored_old._apply_study_state(old_payload)
    assert restored_old.tab_semi.chk_axis_lift.isChecked() is False

    window.close()
    restored.close()
    restored_old.close()


def test_section_panel_import_old_state_uses_default_chapon_length():
    _app()
    panel = SectionCheckPanel()

    panel.import_state({"include_chapon": True, "espesor_chapon": 6.35, "rows": []})

    assert panel.chk_chapon.isChecked() is True
    assert panel.n_chapon_length.value() == 2000.0
    state = panel.export_state()
    assert state["chapon_length_mm"] == 2000.0


def test_section_panel_import_old_lateral_frame_false_structural_forces_true():
    _app()
    panel = SectionCheckPanel()
    panel.import_state({
        "include_bastidor_lateral": True,
        "bastidor_lateral_structural": False,
        "bastidor_lateral_height_mm": 150.0,
    })

    assert panel.chk_bastidor_lateral.isChecked() is True
    assert panel.chk_bastidor_lateral_structural.isChecked() is True
    assert panel.export_state()["bastidor_lateral_structural"] is True


def _configure_solvable_new_study(window: FBDApp):
    tab = window.tab_semi
    window.tabs.setCurrentWidget(tab)
    tab.Lc.setValue(13600.0)
    tab.x_front_or_kp.setValue(1250.0)
    tab.R_front_or_kp.setValue(9400.0)
    tab.Rt.setValue(22200.0)
    tab.chk_show_deflection.setChecked(False)
    tab._add_load_row(load_type="Puntual")
    tab.tbl.item(0, tab.COL_MAG).setText("5000")
    tab.tbl.item(0, tab.COL_POS).setText("2500")
    return tab


def _configure_solvable_bitren_study(window: FBDApp):
    tab = window.tab_bitren
    window.tabs.setCurrentWidget(tab)
    tab.Lc.setValue(13600.0)
    tab.x_front_or_kp.setValue(1250.0)
    tab.R_front_or_kp.setValue(14500.0)
    tab.Rt.setValue(22200.0)
    tab.x_rp2_rel.setValue(900.0)
    tab.Rp2.setValue(13200.0)
    tab.chk_show_deflection.setChecked(False)
    tab._add_load_row(load_type="Puntual")
    tab.tbl.item(0, tab.COL_MAG).setText("5000")
    tab.tbl.item(0, tab.COL_POS).setText("2500")
    return tab


def _point_by_label(points, label: str):
    matches = [p for p in points if p.label == label]
    assert len(matches) == 1
    return matches[0]


def test_unit_tab_single_load_table_parses_all_load_types():
    _app()
    window = FBDApp()
    tab = window.tab_semi
    tab.Lc.setValue(12000.0)
    tab.tbl.setRowCount(0)

    tab._add_load_row(load_type="Puntual")
    tab.tbl.item(0, tab.COL_MAG).setText("500")
    tab.tbl.item(0, tab.COL_POS).setText("1000")

    tab._add_load_row(load_type="Distribuida")
    tab.tbl.item(1, tab.COL_MAG).setText("1200")
    tab.tbl.item(1, tab.COL_POS).setText("4000")
    tab.tbl.item(1, tab.COL_LEN).setText("600")

    tab._add_load_row(load_type="Momento")
    tab.tbl.item(2, tab.COL_MAG).setText("25000")
    tab.tbl.item(2, tab.COL_POS).setText("7000")

    _beam, points, dists, moms, notes = tab.parse_inputs()

    assert notes == []
    assert len(points) == 1
    assert isinstance(points[0], PointForce)
    assert points[0].label == "P1"
    assert points[0].value_user == 500.0
    assert points[0].x_mm == 1000.0
    assert len(dists) == 1
    assert isinstance(dists[0], DistUniform)
    assert dists[0].label == "q1"
    assert dists[0].x0_mm == 3700.0
    assert dists[0].Lq_mm == 600.0
    assert dists[0].q_user == 2.0
    assert len(moms) == 1
    assert isinstance(moms[0], PointMoment)
    assert moms[0].label == "M1"
    assert moms[0].M_user_kgmm == 25000.0
    assert moms[0].x_mm == 7000.0
    window.close()


def test_unit_tab_imports_legacy_split_load_tables_into_single_load_table():
    _app()
    window = FBDApp()
    tab = window.tab_semi
    legacy_state = tab.export_state()
    legacy_state.pop("loads", None)
    legacy_state["tbl_points"] = [["P1", "1000", "500"]]
    legacy_state["tbl_dists"] = [["q1", "3700", "600", "2"]]
    legacy_state["tbl_moms"] = [["M1", "7000", "25000"]]

    restored = FBDApp()
    restored.tab_semi.import_state(legacy_state)
    restored_tab = restored.tab_semi

    assert restored_tab.tbl.rowCount() == 3
    assert restored_tab._type_combo(0).currentText() == "Puntual"
    assert restored_tab.tbl.item(0, restored_tab.COL_MAG).text() == "500"
    assert restored_tab.tbl.item(0, restored_tab.COL_POS).text() == "1000"
    assert restored_tab._type_combo(1).currentText() == "Distribuida"
    assert restored_tab.tbl.item(1, restored_tab.COL_MAG).text() == "1200"
    assert restored_tab.tbl.item(1, restored_tab.COL_POS).text() == "4000"
    assert restored_tab.tbl.item(1, restored_tab.COL_LEN).text() == "600"
    assert restored_tab._type_combo(2).currentText() == "Momento"
    assert restored_tab.tbl.item(2, restored_tab.COL_MAG).text() == "25000"
    assert restored_tab.tbl.item(2, restored_tab.COL_POS).text() == "7000"
    window.close()
    restored.close()


def test_axis_lift_adds_single_first_axle_load_without_moving_tandem():
    _app()
    window = FBDApp()
    tab = _configure_solvable_new_study(window)

    window._solve_for_tab(tab)
    base_cache = tab.get_cache()
    assert base_cache is not None
    base_rp1 = _point_by_label(base_cache.points, "Rp1")
    base_rt = _point_by_label(base_cache.points, "Rt")
    sample_x = float(base_rt.x_mm) - 500.0
    base_v = float(tab.get_diag().eval_V(sample_x))
    base_m = float(tab.get_diag().eval_M(sample_x))
    assert [p for p in base_cache.points if p.label.startswith("P_levante")] == []

    tab.chk_axis_lift.setChecked(True)
    window._replot_active_tab()
    lifted_cache = tab.get_cache()
    assert lifted_cache is not None
    lifted_rp1 = _point_by_label(lifted_cache.points, "Rp1")
    lifted_rt = _point_by_label(lifted_cache.points, "Rt")
    lift_loads = [p for p in lifted_cache.points if p.label == "P_levante_primer_eje"]

    assert len(lift_loads) == 1
    assert lift_loads[0].value_user == 1200.0
    assert lift_loads[0].x_mm == lifted_rt.x_mm - 625.0
    assert lifted_rt.x_mm == base_rt.x_mm
    assert lifted_rt.value_user != base_rt.value_user
    assert (lifted_rp1.value_user, lifted_rt.value_user) != (base_rp1.value_user, base_rt.value_user)
    assert tab.view_mode() == "solved"
    assert float(tab.get_diag().eval_V(sample_x)) != base_v
    assert float(tab.get_diag().eval_M(sample_x)) != base_m
    assert abs(float(tab.get_diag().eval_V(lifted_cache.beam_plot.L_mm))) < 1e-6
    assert abs(float(tab.get_diag().eval_M(lifted_cache.beam_plot.L_mm))) < 1e-3

    tab.chk_axis_lift.setChecked(False)
    window._replot_active_tab()
    disabled_cache = tab.get_cache()
    assert disabled_cache is not None
    assert [p for p in disabled_cache.points if p.label.startswith("P_levante")] == []
    assert float(tab.get_diag().eval_V(sample_x)) == base_v
    assert float(tab.get_diag().eval_M(sample_x)) == base_m

    tab.chk_axis_lift.setChecked(True)
    window._replot_active_tab()
    window._solve_for_tab(tab)
    solved_again = tab.get_cache()
    assert solved_again is not None
    assert len([p for p in solved_again.points if p.label == "P_levante_primer_eje"]) == 1
    window.close()


def test_bitren_axis_lift_keeps_rp2_in_final_equilibrium():
    _app()
    window = FBDApp()
    tab = _configure_solvable_bitren_study(window)

    window._solve_for_tab(tab)
    base_cache = tab.get_cache()
    assert base_cache is not None
    base_rp1 = _point_by_label(base_cache.points, "Rp1")
    base_rp2 = _point_by_label(base_cache.points, "Rp2")
    base_rt = _point_by_label(base_cache.points, "Rt")
    assert base_rt.x_mm > tab.Lc.value()
    sample_x = float(base_rt.x_mm) - tab._first_axle_offset_from_tandem_center() + 100.0
    base_v = float(tab.get_diag().eval_V(sample_x))
    base_m = float(tab.get_diag().eval_M(sample_x))

    tab.chk_axis_lift.setChecked(True)
    window._replot_active_tab()
    lifted_cache = tab.get_cache()
    assert lifted_cache is not None
    lifted_rp1 = _point_by_label(lifted_cache.points, "Rp1")
    lifted_rp2 = _point_by_label(lifted_cache.points, "Rp2")
    lifted_rt = _point_by_label(lifted_cache.points, "Rt")
    lift_load = _point_by_label(lifted_cache.points, "P_levante_primer_eje")

    assert lifted_rp2.x_mm == base_rp2.x_mm
    assert lifted_rp2.value_user == base_rp2.value_user == tab.Rp2.value()
    assert lift_load.value_user == 1200.0
    assert lift_load.x_mm == lifted_rt.x_mm - tab._first_axle_offset_from_tandem_center()
    assert lifted_rt.x_mm == base_rt.x_mm
    assert (lifted_rp1.value_user, lifted_rt.value_user) != (base_rp1.value_user, base_rt.value_user)
    assert float(tab.get_diag().eval_V(sample_x)) != base_v
    assert float(tab.get_diag().eval_M(sample_x)) != base_m
    assert abs(float(tab.get_diag().eval_V(lifted_cache.beam_plot.L_mm))) < 1e-6
    assert abs(float(tab.get_diag().eval_M(lifted_cache.beam_plot.L_mm))) < 1e-3
    window.close()


def test_axis_lift_directional_load_uses_directional_position():
    _app()
    window = FBDApp()
    tab = _configure_solvable_new_study(window)
    tab.cmb_config.setCurrentText("1 + 2 ejes — Rd 9200 (offset 3075) + Rt 15800")
    tab.Rd.setValue(9200.0)
    tab.dir_offset.setValue(3075.0)
    tab.chk_axis_lift.setChecked(True)

    window._solve_for_tab(tab)

    cache = tab.get_cache()
    assert cache is not None
    rt = _point_by_label(cache.points, "Rt")
    lift_load = _point_by_label(cache.points, "P_levante_direccional")
    assert "Levantar direccional" in tab.chk_axis_lift.text()
    assert lift_load.value_user == 1300.0
    assert tab.directional_offset_for_solver() == tab.dir_offset.value() + 625.0
    assert lift_load.x_mm == rt.x_mm - tab.directional_offset_for_solver()
    assert [p for p in cache.points if p.label == "Rd"] == []
    assert abs(float(tab.get_diag().eval_V(cache.beam_plot.L_mm))) < 1e-6
    assert abs(float(tab.get_diag().eval_M(cache.beam_plot.L_mm))) < 1e-3
    window.close()


def test_axis_lift_treats_1_1_1_as_directional_lift():
    _app()
    window = FBDApp()
    tab = window.tab_semi
    tab.cmb_config.setCurrentText("1 + 1 + 1 ejes — Rd 9200 (offset 2450) + Rt 18800")
    tab._apply_config_defaults()

    assert tab.chk_axis_lift.isEnabled() is True
    assert "Levantar direccional" in tab.chk_axis_lift.text()
    assert tab.Rd.value() == 9200.0
    assert tab.Rt.value() == 18800.0
    assert tab.Rd.value() + tab.Rt.value() == 28000.0
    assert tab.dir_offset.value() == 2450.0
    window.close()


def test_unit_tab_accepts_positive_load_positions_beyond_visible_beam(tmp_path):
    _app()
    window = FBDApp()
    tab = window.tab_semi
    tab.Lc.setValue(12000.0)
    tab.chk_show_deflection.setChecked(False)
    tab.tbl.setRowCount(0)

    tab._add_load_row(load_type="Puntual")
    tab.tbl.item(0, tab.COL_MAG).setText("500")
    tab.tbl.item(0, tab.COL_POS).setText("12500")

    tab._add_load_row(load_type="Puntual")
    tab.tbl.item(1, tab.COL_MAG).setText("600")
    tab.tbl.item(1, tab.COL_POS).setText("13000")

    tab._add_load_row(load_type="Momento")
    tab.tbl.item(2, tab.COL_MAG).setText("25000")
    tab.tbl.item(2, tab.COL_POS).setText("12750")

    tab._add_load_row(load_type="Momento")
    tab.tbl.item(3, tab.COL_MAG).setText("-15000")
    tab.tbl.item(3, tab.COL_POS).setText("13250")

    tab._add_load_row(load_type="Distribuida")
    tab.tbl.item(4, tab.COL_MAG).setText("1200")
    tab.tbl.item(4, tab.COL_POS).setText("12300")
    tab.tbl.item(4, tab.COL_LEN).setText("600")

    points, dists, moms, errors = tab._build_loads_from_table()
    assert errors == []
    assert [p.x_mm for p in points] == [12500.0, 13000.0]
    assert [m.x_mm for m in moms] == [12750.0, 13250.0]
    assert len(dists) == 1
    assert dists[0].x0_mm == 12000.0
    assert dists[0].Lq_mm == 600.0

    tab.tbl.item(4, tab.COL_LEN).setText("-1")
    _, _, _, invalid_errors = tab._build_loads_from_table()
    assert any("longitud > 0" in error for error in invalid_errors)
    tab.tbl.item(4, tab.COL_LEN).setText("600")

    window._solve_for_tab(tab)
    cache = tab.get_cache()
    assert cache is not None
    assert tab.get_diag() is not None

    study_path = tmp_path / "loads_outside_beam.sbeam"
    save_study_file(study_path, window._export_study_state())
    restored = FBDApp()
    restored._apply_study_state(load_study_file(study_path))
    restored_tab = restored.tab_semi
    assert restored_tab.tbl.item(0, restored_tab.COL_POS).text() == "12500"
    assert restored_tab.tbl.item(1, restored_tab.COL_POS).text() == "13000"
    assert restored_tab.tbl.item(2, restored_tab.COL_POS).text() == "12750"
    assert restored_tab.tbl.item(3, restored_tab.COL_POS).text() == "13250"
    assert restored_tab.tbl.item(4, restored_tab.COL_POS).text() == "12300"
    window.close()
    restored.close()


def test_unit_tab_rejects_negative_load_position():
    _app()
    window = FBDApp()
    tab = window.tab_semi
    tab.Lc.setValue(12000.0)
    tab.tbl.setRowCount(0)

    tab._add_load_row(load_type="Puntual")
    tab.tbl.item(0, tab.COL_MAG).setText("500")
    tab.tbl.item(0, tab.COL_POS).setText("-500")

    _points, _dists, _moms, errors = tab._build_loads_from_table()
    assert any("posición no puede ser negativa" in error for error in errors)
    window.close()


def test_reactions_tab_accepts_positions_beyond_reference_length():
    _app()
    window = FBDApp()
    tab = window.tab_reactions
    window.tabs.setCurrentWidget(tab)
    tab.mode.setCurrentIndex(0)
    tab.L.setValue(12000.0)
    tab.x_a.setValue(1250.0)
    tab.x_b.setValue(13000.0)
    tab.tbl.setRowCount(0)

    tab._add_load_row(load_type="Puntual")
    tab.tbl.item(0, tab.COL_MAG).setText("500")
    tab.tbl.item(0, tab.COL_POS).setText("12500")

    tab._add_load_row(load_type="Momento")
    tab.tbl.item(1, tab.COL_MAG).setText("25000")
    tab.tbl.item(1, tab.COL_POS).setText("12750")

    tab._add_load_row(load_type="Distribuida")
    tab.tbl.item(2, tab.COL_MAG).setText("1200")
    tab.tbl.item(2, tab.COL_POS).setText("12300")
    tab.tbl.item(2, tab.COL_LEN).setText("600")

    loads, errors, _ = tab._build_loads()
    assert errors == []
    assert [load.x_mm for load in loads if isinstance(load, PointForce)] == [12500.0]
    assert [load.x_mm for load in loads if isinstance(load, PointMoment)] == [12750.0]
    assert [load.x0_mm for load in loads if isinstance(load, DistUniform)] == [12000.0]

    tab.recompute_now()
    assert tab._last_result is not None
    assert "0 <= " not in tab.note_label.text()
    assert "entre 0 y L" not in tab.note_label.text()
    window.close()


def test_reactions_tab_accepts_three_support_axis_beyond_reference_length():
    _app()
    window = FBDApp()
    tab = window.tab_reactions
    window.tabs.setCurrentWidget(tab)
    tab.mode.setCurrentIndex(1)
    tab.L.setValue(12000.0)
    tab.x_k.setValue(1250.0)
    tab.x_t.setValue(13200.0)
    tab.offset.setValue(3700.0)
    tab.x_t_min.setValue(5000.0)
    tab.x_t_max.setValue(14000.0)
    tab.tbl.setRowCount(0)

    tab._add_load_row(load_type="Puntual")
    tab.tbl.item(0, tab.COL_MAG).setText("500")
    tab.tbl.item(0, tab.COL_POS).setText("12500")

    tab.recompute_now()
    assert tab._last_result is not None
    assert "0 <= " not in tab.note_label.text()
    assert "entre 0 y L" not in tab.note_label.text()
    window.close()


def test_reactions_tab_rejects_negative_position_and_invalid_dist_length():
    _app()
    window = FBDApp()
    tab = window.tab_reactions
    tab.L.setValue(12000.0)
    tab.tbl.setRowCount(0)

    tab._add_load_row(load_type="Puntual")
    tab.tbl.item(0, tab.COL_MAG).setText("500")
    tab.tbl.item(0, tab.COL_POS).setText("-1")

    tab._add_load_row(load_type="Distribuida")
    tab.tbl.item(1, tab.COL_MAG).setText("1200")
    tab.tbl.item(1, tab.COL_POS).setText("1000")
    tab.tbl.item(1, tab.COL_LEN).setText("-600")

    _loads, errors, _ = tab._build_loads()
    assert any("posición no puede ser negativa" in error for error in errors)
    assert any("longitud > 0" in error for error in errors)
    assert all("entre 0 y L" not in error for error in errors)
    window.close()


def test_unit_tabs_expose_load_mode_selector_without_changing_reactions_tab():
    _app()
    window = FBDApp()

    for tab in (window.tab_acoplado, window.tab_semi, window.tab_bitren):
        assert tab.cmb_load_mode.itemText(0) == "Carga distribuida equivalente"
        assert tab.cmb_load_mode.itemText(1) == "Cargas reales"
        assert tab.load_mode() == "Carga distribuida equivalente"

    assert not hasattr(window.tab_reactions, "cmb_load_mode")
    window.close()


def test_real_load_changes_do_not_add_mode_flow_to_reactions_tab():
    _app()
    window = FBDApp()

    assert not hasattr(window.tab_reactions, "cmb_load_mode")
    assert "Buscar (cumplir" in window.tab_reactions.btn_search.text()
    assert window.tab_reactions.btn_search_best.text() == "Buscar mejor margen"
    window.close()


def test_load_mode_roundtrip_and_legacy_default(tmp_path):
    _app()
    window = FBDApp()
    window.tab_acoplado.cmb_load_mode.setCurrentText("Carga distribuida equivalente")
    window.tab_semi.cmb_load_mode.setCurrentText("Cargas reales")
    window.tab_bitren.cmb_load_mode.setCurrentText("Cargas reales")

    payload = window._export_study_state()
    study_path = tmp_path / "load_modes.sbeam"
    save_study_file(study_path, payload)

    restored = FBDApp()
    restored._apply_study_state(load_study_file(study_path))
    assert restored.tab_acoplado.load_mode() == "Carga distribuida equivalente"
    assert restored.tab_semi.load_mode() == "Cargas reales"
    assert restored.tab_bitren.load_mode() == "Cargas reales"

    legacy_payload = dict(payload)
    legacy_tabs = {
        key: dict(value)
        for key, value in payload["tabs"].items()
    }
    for key in ("acoplado", "semirremolque", "bitren"):
        legacy_tabs[key].pop("load_mode", None)
    legacy_payload["tabs"] = legacy_tabs

    legacy = FBDApp()
    legacy._apply_study_state(legacy_payload)
    assert legacy.tab_acoplado.load_mode() == "Carga distribuida equivalente"
    assert legacy.tab_semi.load_mode() == "Carga distribuida equivalente"
    assert legacy.tab_bitren.load_mode() == "Carga distribuida equivalente"
    window.close()
    restored.close()
    legacy.close()


def test_real_load_mode_uses_table_loads_without_equivalent_q():
    _app()
    window = FBDApp()
    tab = _configure_solvable_new_study(window)
    tab.cmb_load_mode.setCurrentText("Cargas reales")
    tab.tbl.setRowCount(0)

    tab._add_load_row(load_type="Puntual")
    tab.tbl.item(0, tab.COL_MAG).setText("5000")
    tab.tbl.item(0, tab.COL_POS).setText("2500")

    tab._add_load_row(load_type="Distribuida")
    tab.tbl.item(1, tab.COL_MAG).setText("1200")
    tab.tbl.item(1, tab.COL_POS).setText("6000")
    tab.tbl.item(1, tab.COL_LEN).setText("600")

    tab._add_load_row(load_type="Momento")
    tab.tbl.item(2, tab.COL_MAG).setText("25000")
    tab.tbl.item(2, tab.COL_POS).setText("7000")

    window._solve_for_tab(tab)
    cache = tab.get_cache()
    assert cache is not None
    assert tab.view_mode() == "solved"
    assert tab.get_diag() is not None
    assert "Modo de carga = Cargas reales" in cache.note_text
    assert "q equivalente automática = no aplicada" in cache.note_text
    assert not any(load.label == "q" for load in cache.dists)
    assert [load.label for load in cache.dists] == ["q1"]
    assert _point_by_label(cache.points, "P1").x_mm == 2500.0
    assert cache.moms[0].label == "M1"
    assert not tab.lbl_reaction_limits.isHidden()
    assert "Uso de límites admisibles" in tab.lbl_reaction_limits.text()
    assert "Rp1:" in tab.lbl_reaction_limits.text()
    assert "Rt:" in tab.lbl_reaction_limits.text()
    assert "% del límite" in tab.lbl_reaction_limits.text()
    window.close()


def test_real_load_mode_does_not_call_equivalent_geometry_solver(monkeypatch):
    import semi_beam.ui.main_window as main_window

    _app()
    window = FBDApp()
    tab = _configure_solvable_new_study(window)
    tab.cmb_load_mode.setCurrentText("Cargas reales")

    def fail_solve_equilibrium(_case):
        raise AssertionError("solve_equilibrium must not run in real load mode")

    monkeypatch.setattr(main_window, "solve_equilibrium", fail_solve_equilibrium)

    window._solve_for_tab(tab)

    cache = tab.get_cache()
    assert cache is not None
    assert "Busqueda x_t" in cache.note_text
    assert "q equivalente autom" in cache.note_text
    assert "geometr" not in cache.note_text.lower()
    window.close()


def test_real_load_manual_fixed_tandem_candidate_matches_two_support_equilibrium():
    _app()
    window = FBDApp()
    tab = window.tab_semi
    tab.R_front_or_kp.setValue(7000.0)
    tab.Rt.setValue(6000.0)
    case = BeamCase(
        beam=Beam(L_mm=12000.0),
        point_forces=[
            PointForce(label="P1", x_mm=3000.0, value_user=6000.0),
            PointForce(label="P2", x_mm=7000.0, value_user=4000.0),
        ],
        dist_loads=[],
        moments=[],
        kingpin=FixedSupport(label="Rp1", x_mm=0.0, reaction_user=0.0),
        tandem=TandemSupport(label="Rt", reaction_user=0.0),
        unknown_uniform=UnknownUniformLoad(label="q", span_start_mm=0.0, span_len_mm=12000.0),
    )

    candidate = window._evaluate_real_load_candidate(
        tab=tab,
        case=case,
        beam_L_mm=12000.0,
        x_t=8000.0,
    )
    by_label = {point.label: point for point in candidate.support_points}

    assert abs(by_label["Rp1"].value_user - 4250.0) < 1e-6
    assert abs(by_label["Rt"].value_user - 5750.0) < 1e-6
    assert candidate.feasible is True
    window.close()


def test_real_load_tandem_search_finds_admissible_position_by_documented_criterion():
    _app()
    window = FBDApp()
    tab = window.tab_semi
    tab.Lc.setValue(12000.0)
    tab.x_front_or_kp.setValue(0.0)
    tab.R_front_or_kp.setValue(7000.0)
    tab.Rt.setValue(6000.0)
    case = BeamCase(
        beam=Beam(L_mm=12000.0),
        point_forces=[
            PointForce(label="P1", x_mm=3000.0, value_user=6000.0),
            PointForce(label="P2", x_mm=7000.0, value_user=4000.0),
        ],
        dist_loads=[],
        moments=[],
        kingpin=FixedSupport(label="Rp1", x_mm=0.0, reaction_user=0.0),
        tandem=TandemSupport(label="Rt", reaction_user=0.0),
        unknown_uniform=UnknownUniformLoad(label="q", span_start_mm=0.0, span_len_mm=12000.0),
    )

    at_7000 = window._evaluate_real_load_candidate(tab=tab, case=case, beam_L_mm=12000.0, x_t=7000.0)
    at_8000 = window._evaluate_real_load_candidate(tab=tab, case=case, beam_L_mm=12000.0, x_t=8000.0)
    search = window._search_real_load_tandem_position(tab=tab, case=case, Lc=12000.0)

    assert at_7000.feasible is False
    assert at_8000.feasible is True
    assert search.feasible_count > 0
    assert search.candidate.feasible is True
    assert search.candidate.max_usage_pct <= at_8000.max_usage_pct
    window.close()


def test_real_load_search_range_is_lc_half_to_lc_plus_3000_for_all_unit_tabs():
    _app()
    window = FBDApp()

    for tab in (window.tab_acoplado, window.tab_semi, window.tab_bitren):
        tab.Lc.setValue(12400.0)
        tab.x_front_or_kp.setValue(0.0)
        hitch = None
        if tab.is_bitren:
            tab.x_rp2_rel.setValue(1800.0)
            tab.Rp2.setValue(13200.0)
            hitch = FixedSupport(label="Rp2", x_mm=14200.0, reaction_user=13200.0)
        case = BeamCase(
            beam=Beam(L_mm=12400.0),
            point_forces=[],
            dist_loads=[],
            moments=[],
            kingpin=FixedSupport(label="Rp1", x_mm=0.0, reaction_user=0.0),
            tandem=TandemSupport(label="Rt", reaction_user=0.0),
            hitch=hitch,
            unknown_uniform=UnknownUniformLoad(label="q", span_start_mm=0.0, span_len_mm=12400.0),
        )

        lo, hi, step = window._real_load_search_range(tab=tab, case=case, Lc=12400.0)
        values = window._iter_real_load_xt_candidates(lo, hi, step)

        assert lo == 6200.0
        assert hi == 15400.0
        assert 12950.0 in values
        assert hi > tab.Lc.value()

    window.close()


def test_bitren_real_load_search_allows_tandem_before_rp2_and_keeps_rp2():
    _app()
    window = FBDApp()
    tab = window.tab_bitren
    tab.cmb_load_mode.setCurrentText("Cargas reales")
    tab.Lc.setValue(12400.0)
    tab.x_front_or_kp.setValue(0.0)
    tab.R_front_or_kp.setValue(20000.0)
    tab.Rt.setValue(25000.0)
    tab.x_rp2_rel.setValue(1800.0)
    tab.Rp2.setValue(13200.0)
    case = BeamCase(
        beam=Beam(L_mm=12400.0),
        point_forces=[
            PointForce(label="P1", x_mm=3000.0, value_user=6000.0),
            PointForce(label="P2", x_mm=7000.0, value_user=4000.0),
        ],
        dist_loads=[],
        moments=[],
        kingpin=FixedSupport(label="Rp1", x_mm=0.0, reaction_user=0.0),
        tandem=TandemSupport(label="Rt", reaction_user=0.0),
        hitch=FixedSupport(label="Rp2", x_mm=14200.0, reaction_user=13200.0),
        unknown_uniform=UnknownUniformLoad(label="q", span_start_mm=0.0, span_len_mm=12400.0),
    )

    lo, hi, step = window._real_load_search_range(tab=tab, case=case, Lc=12400.0)
    values = window._iter_real_load_xt_candidates(lo, hi, step)
    candidate = window._evaluate_real_load_candidate(
        tab=tab,
        case=case,
        beam_L_mm=15400.0,
        x_t=12950.0,
    )
    by_label = {point.label: point for point in candidate.support_points}

    assert 12950.0 in values
    assert 12950.0 < 14200.0
    assert by_label["Rt"].x_mm == 12950.0
    assert by_label["Rp2"].x_mm == 14200.0
    assert by_label["Rp2"].value_user == 13200.0
    window.close()


def test_real_load_tandem_search_reports_no_admissible_position():
    _app()
    window = FBDApp()
    tab = window.tab_semi
    tab.cmb_load_mode.setCurrentText("Cargas reales")
    tab.Lc.setValue(12000.0)
    tab.x_front_or_kp.setValue(0.0)
    tab.R_front_or_kp.setValue(1000.0)
    tab.Rt.setValue(1000.0)
    tab.tbl.setRowCount(0)
    tab._add_load_row(load_type="Puntual")
    tab.tbl.item(0, tab.COL_MAG).setText("6000")
    tab.tbl.item(0, tab.COL_POS).setText("3000")
    tab._add_load_row(load_type="Puntual")
    tab.tbl.item(1, tab.COL_MAG).setText("4000")
    tab.tbl.item(1, tab.COL_POS).setText("7000")

    window._solve_for_tab(tab)

    cache = tab.get_cache()
    assert cache is not None
    assert "no se encontro posicion admisible" in cache.note_text
    assert "excedido" in cache.note_text
    window.close()


def test_real_load_reaction_fields_explain_limit_semantics_without_breaking_equivalent_mode():
    _app()
    window = FBDApp()
    tab = window.tab_semi

    assert "limite admisible" in tab.R_front_or_kp.toolTip()
    assert "limite admisible" in tab.Rt.toolTip()
    assert "carga equivalente" in tab.R_front_or_kp.toolTip()
    assert tab.load_mode() == "Carga distribuida equivalente"
    window.close()


def test_equivalent_load_mode_keeps_reaction_limit_summary_hidden():
    _app()
    window = FBDApp()
    tab = _configure_solvable_new_study(window)
    tab.cmb_load_mode.setCurrentText("Carga distribuida equivalente")

    window._solve_for_tab(tab)

    cache = tab.get_cache()
    assert cache is not None
    assert "q calculada" in cache.note_text
    assert "Uso de límites admisibles" not in cache.note_text
    assert not tab.lbl_reaction_limits.isVisible()
    assert tab.lbl_reaction_limits.text() == ""
    window.close()


def test_real_load_reaction_limit_usage_formula_and_status_colors():
    _app()
    window = FBDApp()
    tab = window.tab_acoplado
    tab.R_front_or_kp.setValue(10000.0)
    tab.Rt.setValue(10000.0)
    points = [
        PointForce(label="Rp1", x_mm=0.0, value_user=10860.0),
        PointForce(label="Rt", x_mm=1000.0, value_user=3940.0),
    ]

    usages = window._reaction_limit_usages(tab, points)
    by_label = {usage.label: usage for usage in usages}

    assert abs(by_label["Rp1"].percent - 108.6) < 1e-9
    assert by_label["Rp1"].exceeded is True
    assert by_label["Rp1"].status == "excedido"
    assert abs(by_label["Rt"].percent - 39.4) < 1e-9
    assert by_label["Rt"].exceeded is False
    assert by_label["Rt"].status == "admisible"

    tab.set_reaction_limit_summary(usages)
    summary = tab.lbl_reaction_limits.text()
    assert "#B00020" in summary
    assert "#0A7F2E" in summary
    assert "108.6 % del límite" in summary
    assert "39.4 % del límite" in summary
    window.close()


def test_bitren_real_load_mode_keeps_rp2():
    _app()
    window = FBDApp()
    tab = _configure_solvable_bitren_study(window)
    tab.cmb_load_mode.setCurrentText("Cargas reales")

    window._solve_for_tab(tab)
    cache = tab.get_cache()
    assert cache is not None
    rp2 = _point_by_label(cache.points, "Rp2")
    assert rp2.value_user == tab.Rp2.value()
    assert "q equivalente automática = no aplicada" in cache.note_text
    assert not any(load.label == "q" for load in cache.dists)
    assert tab.get_diag() is not None
    assert "Uso de límites admisibles" in cache.note_text
    assert "Rp1:" in cache.note_text
    assert "límite 14500 kg" in cache.note_text
    assert "Rt:" in cache.note_text
    assert "límite 22200 kg" in cache.note_text
    assert "Rp2:" in cache.note_text
    assert "límite 13200 kg" in cache.note_text
    assert not tab.lbl_reaction_limits.isHidden()
    window.close()


def test_directional_offset_is_entered_from_second_axis_and_validated():
    _app()
    window = FBDApp()
    tab = _configure_solvable_new_study(window)
    tab.cmb_config.setCurrentText("1 + 2 ejes — Rd 9200 (offset 3075) + Rt 15800")
    tab._apply_config_defaults()

    assert tab.dir_offset.minimum() == DIRECTIONAL_OFFSET_MIN_MM
    assert tab.dir_offset.maximum() == DIRECTIONAL_OFFSET_MAX_MM
    assert "segundo eje" in tab.dir_offset.toolTip()

    for value in (2400.0, 2450.0, 2475.0, 4000.0):
        tab.dir_offset.setValue(value)
        assert tab._validate_required_inputs() == []

    for invalid in ("0", "1000", "4500"):
        tab.dir_offset.lineEdit().setText(invalid)
        assert any("2400 y 4000" in error for error in tab._validate_required_inputs())

    tab.dir_offset.setValue(3075.0)
    assert tab.directional_offset_from_second_axis() == 3075.0
    assert tab.directional_offset_for_solver() == 3700.0
    window.close()


def test_semi_1_1_1_solves_with_directional_lift_and_fixed_tandem():
    _app()
    window = FBDApp()
    tab = _configure_solvable_new_study(window)
    tab.cmb_config.setCurrentText("1 + 1 + 1 ejes — Rd 9200 (offset 2450) + Rt 18800")
    tab._apply_config_defaults()

    window._solve_for_tab(tab)
    base_cache = tab.get_cache()
    assert base_cache is not None
    base_rt = _point_by_label(base_cache.points, "Rt")

    tab.chk_axis_lift.setChecked(True)
    window._replot_active_tab()
    lifted_cache = tab.get_cache()
    assert lifted_cache is not None
    lifted_rt = _point_by_label(lifted_cache.points, "Rt")
    lift_load = _point_by_label(lifted_cache.points, "P_levante_direccional")

    assert tab.Rd.value() + tab.Rt.value() == 28000.0
    assert tab.dir_offset.value() == 2450.0
    assert lift_load.value_user == 1300.0
    assert tab.directional_offset_for_solver() == 3675.0
    assert lift_load.x_mm == lifted_rt.x_mm - tab.directional_offset_for_solver()
    assert lifted_rt.x_mm == base_rt.x_mm
    assert lifted_rt.value_user != base_rt.value_user
    assert abs(float(tab.get_diag().eval_V(lifted_cache.beam_plot.L_mm))) < 1e-6
    assert abs(float(tab.get_diag().eval_M(lifted_cache.beam_plot.L_mm))) < 1e-3
    window.close()


def test_unit_tabs_have_unified_load_add_controls_and_fixed_tab_bar():
    _app()
    window = FBDApp()

    assert isinstance(window.tab_acoplado.content_scroll, QScrollArea)
    assert isinstance(window.tab_semi.content_scroll, QScrollArea)
    assert isinstance(window.tab_bitren.content_scroll, QScrollArea)
    assert isinstance(window.tab_reactions.content_scroll, QScrollArea)
    assert not isinstance(window.tabs.parentWidget(), QScrollArea)

    for tab in (window.tab_acoplado, window.tab_semi, window.tab_bitren):
        sections = [button for button in tab.findChildren(QToolButton) if button.text() == "Cargas"]
        assert len(sections) == 1
        assert tab.btn_add_load.text() == "Agregar carga"
        assert tab.btn_del_load.text() == "Eliminar seleccionadas"
        assert not hasattr(tab, "cmb_add_load_type")
        assert not hasattr(tab, "btn_add_p")
        assert not hasattr(tab, "btn_add_q")
        assert not hasattr(tab, "btn_add_m")
        assert not hasattr(tab, "tbl_points")
        assert not hasattr(tab, "tbl_dists")
        assert not hasattr(tab, "tbl_moms")

        headers = [
            tab.tbl.horizontalHeaderItem(col).text()
            for col in range(tab.tbl.columnCount())
        ]
        assert headers == ["Tipo", "Magnitud", "Posición / centro [mm]", "Longitud [mm]"]

        for load_type in ("Puntual", "Distribuida", "Momento"):
            before = tab.tbl.rowCount()
            tab._add_load_row(load_type=load_type)
            assert tab.tbl.rowCount() == before + 1
            assert tab._type_combo(before).currentText() == load_type

    assert window.tab_reactions.btn_add.text() == "Agregar carga"
    window.close()


def test_unit_tabs_render_single_general_load_table_without_old_collapsibles():
    app = _app()
    window = FBDApp()
    window.resize(1400, 900)
    window.show()
    app.processEvents()

    forbidden_sections = {
        "Fuerzas puntuales conocidas (P1, P2, ...)",
        "Cargas distribuidas conocidas (kg/mm)",
        "Momentos puntuales (kg·mm, CCW+)",
    }
    general_headers = ("Tipo", "Magnitud", "Posición / centro [mm]", "Longitud [mm]")

    for tab in (window.tab_acoplado, window.tab_semi, window.tab_bitren):
        window.tabs.setCurrentWidget(tab)
        app.processEvents()

        visible_labels = {label.text() for label in tab.findChildren(QLabel) if label.isVisible()}
        visible_buttons = [button for button in tab.findChildren(QToolButton) if button.isVisible()]
        visible_sections = {button.text() for button in visible_buttons}
        assert "Cargas" in visible_sections
        assert visible_sections.isdisjoint(forbidden_sections)
        assert "Bastidor lateral estructural" not in visible_labels
        assert "Deformada" not in visible_sections
        assert "Mostrar deformada" not in visible_labels

        load_button = next(button for button in visible_buttons if button.text() == "Cargas")
        load_button.click()
        app.processEvents()
        visible_labels = {label.text() for label in tab.findChildren(QLabel) if label.isVisible()}
        assert "Magnitud: kg para puntuales y distribuidas (P total), kg·mm para momentos." in visible_labels

        visible_headers = []
        for table in tab.findChildren(QTableWidget):
            if not table.isVisible():
                continue
            visible_headers.append(tuple(
                table.horizontalHeaderItem(col).text()
                for col in range(table.columnCount())
            ))

        assert visible_headers.count(general_headers) == 1
        assert ("label", "x_mm", "valor_kg") not in visible_headers
        assert ("label", "x0_mm", "Lq_mm", "q_kg_per_mm") not in visible_headers
        assert ("label", "x_mm", "M_kgmm") not in visible_headers

    bitren_labels = {
        label.text()
        for label in window.tab_bitren.findChildren(QLabel)
        if label.isVisible()
    }
    bitren_buttons = {
        button.text()
        for button in window.tab_bitren.findChildren(QToolButton)
        if button.isVisible()
    }
    assert not any("direccional" in text.lower() for text in bitren_labels | bitren_buttons)

    window.close()


def test_about_action_is_single_button_at_top_bar_right(monkeypatch):
    app = _app()
    window = FBDApp()
    window.show()
    app.processEvents()

    assert [action.text() for action in window.menuBar().actions()] == []

    about_buttons = [
        button
        for button in window.findChildren(QPushButton)
        if button.text() == "Acerca de Calculeitor"
    ]
    assert about_buttons == [window.btn_about]
    assert window.btn_about.parentWidget() is window.btn_export_memoria_docx.parentWidget()
    assert window.btn_about.geometry().x() > window.btn_export_memoria_docx.geometry().x()
    assert window.btn_about.geometry().right() > window.btn_export_memoria_docx.geometry().right()

    tab_texts = [window.tabs.tabText(index) for index in range(window.tabs.count())]
    assert tab_texts == ["Acoplado", "Semirremolque", "Bitren", "Cálculo y verificación"]
    assert not any(
        button.text() == "Ayuda"
        for button in window.findChildren(QPushButton)
    )

    calls = []

    def fake_about(parent, title, text):
        calls.append((parent, title, text))

    monkeypatch.setattr(QMessageBox, "about", fake_about)
    window.btn_about.click()

    assert len(calls) == 1
    assert calls[0][0] is window
    assert calls[0][1] == "Acerca de Calculeitor"
    assert "Calculeitor" in calls[0][2]
    window.close()


def test_axis_lift_disabled_for_non_applicable_configuration():
    _app()
    window = FBDApp()
    tab = window.tab_acoplado
    tab.cmb_config.setCurrentText("4 ejes conv — 15800 / 15800")
    tab.chk_axis_lift.setChecked(True)
    tab._apply_config_defaults()

    assert tab.chk_axis_lift.isEnabled() is False
    assert tab.chk_axis_lift.isChecked() is False
    window.close()


def test_new_study_solve_refreshes_section_check_results_without_import_state():
    _app()
    window = FBDApp()
    tab = _configure_solvable_new_study(window)
    panel = tab.section_panel
    panel.tbl.item(0, panel.COL_X).setText("1000")
    panel.tbl.item(0, panel.COL_HWEB).setText("450")

    window._solve_for_tab(tab)

    fs_text = panel.tbl.item(0, panel.COL_FS).text()
    assert panel.tbl.item(0, panel.COL_M).text()
    assert fs_text
    assert "ERR" not in fs_text
    assert float(fs_text) > 0.0
    window.close()


def test_new_study_solve_refreshes_double_web_section_results_without_import_state():
    _app()
    window = FBDApp()
    tab = _configure_solvable_new_study(window)
    panel = tab.section_panel
    panel._double_web_widgets[0].setChecked(True)
    panel.tbl.item(0, panel.COL_DOUBLE_WEB_OFFSET).setText("20")
    panel.tbl.item(0, panel.COL_X).setText("1000")
    panel.tbl.item(0, panel.COL_HWEB).setText("450")

    window._solve_for_tab(tab)

    fs_text = panel.tbl.item(0, panel.COL_FS).text()
    assert panel.tbl.item(0, panel.COL_M).text()
    assert fs_text
    assert "ERR" not in fs_text
    assert float(fs_text) > 0.0
    window.close()


def test_window_title_shows_current_study_filename(tmp_path):
    _app()
    window = FBDApp()

    assert window.windowTitle() == "Calculeitor - Acoplado / Semirremolque / Bitren"

    window._current_study_path = "C:/tmp/Forestal.sbeam"
    window._update_window_title()

    assert "Forestal" in window.windowTitle()
    assert ".sbeam" not in window.windowTitle()
    assert window.windowTitle() == "Calculeitor - Acoplado / Semirremolque / Bitren — Forestal"
